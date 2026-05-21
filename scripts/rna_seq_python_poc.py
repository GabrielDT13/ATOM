#!/usr/bin/env python3
from __future__ import annotations

import argparse
import base64
import json
import logging
import math
import re
import sys
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any
from zipfile import ZIP_DEFLATED, ZipFile


LOGGER = logging.getLogger("rna_seq_python")

FIGURE_METADATA: dict[str, dict[str, str]] = {
    "pca": {
        "label": "PCA de muestras",
        "caption": "Análisis de componentes principales para inspeccionar separación entre grupos y consistencia entre réplicas.",
        "group": "Control de calidad",
    },
    "sample_distance_heatmap": {
        "label": "Heatmap de distancia entre muestras",
        "caption": "Clustering jerárquico de muestras basado en perfiles de expresión normalizados.",
        "group": "Control de calidad",
    },
    "library_size_barplot": {
        "label": "Barplot de tamaño de librería",
        "caption": "Resumen de profundidad de secuenciación entre muestras seleccionadas antes de la interpretación biológica.",
        "group": "Control de calidad",
    },
    "normalized_expression_boxplot": {
        "label": "Boxplot de expresión normalizada",
        "caption": "Distribución de valores de expresión normalizados entre muestras tras el escalado.",
        "group": "Control de calidad",
    },
    "volcano_plot": {
        "label": "Volcano plot",
        "caption": "Panorama global de expresión diferencial con umbrales de significación y cambio de expresión resaltados.",
        "group": "Expresión diferencial",
    },
    "ma_plot": {
        "label": "MA plot",
        "caption": "Relación entre expresión media y log2 fold change para inspeccionar cambios dependientes de intensidad.",
        "group": "Expresión diferencial",
    },
    "heatmap": {
        "label": "Heatmap de genes significativos",
        "caption": "Heatmap de expresión estandarizada para genes significativos destacados entre muestras seleccionadas.",
        "group": "Expresión diferencial",
    },
    "top_significant_genes_barplot": {
        "label": "Barplot de genes significativos",
        "caption": "Comparación compacta de genes más sobreexpresados e infraexpresados según log2 fold change.",
        "group": "Expresión diferencial",
    },
    "pvalue_histogram": {
        "label": "Histograma de adjusted p-value",
        "caption": "Distribución de significación entre genes evaluados para inspeccionar fuerza global de señal.",
        "group": "Expresión diferencial",
    },
    "gsea_prerank_dotplot": {
        "label": "Dot plot GSEA prerankeado",
        "caption": "Términos enriquecidos principales ordenados por significación y normalized enrichment score.",
        "group": "Enriquecimiento funcional",
    },
    "enrichr_up_dotplot": {
        "label": "Dot plot de enriquecimiento sobreexpresado",
        "caption": "Términos funcionales enriquecidos principales entre genes sobreexpresados.",
        "group": "Enriquecimiento funcional",
    },
    "enrichr_down_dotplot": {
        "label": "Dot plot de enriquecimiento infraexpresado",
        "caption": "Términos funcionales enriquecidos principales entre genes infraexpresados.",
        "group": "Enriquecimiento funcional",
    },
}


@dataclass
class PipelineInputs:
    counts: Any
    metadata: Any
    sample_column: str


@dataclass
class DifferentialExpressionOutputs:
    counts: Any
    metadata: Any
    normalized_counts: Any
    results: Any
    significant: Any
    upregulated: Any
    downregulated: Any
    ranked_scores: Any
    summary: dict[str, Any]


@dataclass
class ReportBranding:
    author: str
    entity_name: str
    entity_logo_path: str


def _import_dependency(module_name: str, package_name: str | None = None) -> Any:
    try:
        return __import__(module_name, fromlist=["*"])
    except ImportError as exc:
        package = package_name or module_name
        raise RuntimeError(
            f"Missing dependency '{package}'. Install it before running this script."
        ) from exc


def _configure_logging(verbose: bool = False) -> None:
    logging.basicConfig(
        level=logging.DEBUG if verbose else logging.INFO,
        format="%(asctime)s | %(levelname)s | %(message)s",
    )


def _read_tabular_file(file_path: Path) -> Any:
    pd = _import_dependency("pandas")

    suffix = file_path.suffix.lower()
    if suffix in {".xlsx", ".xls"}:
        return pd.read_excel(file_path)
    if suffix in {".tsv", ".txt"}:
        return pd.read_csv(file_path, sep="\t")
    if suffix == ".csv":
        return pd.read_csv(file_path)

    return pd.read_csv(file_path, sep=None, engine="python")


def _prepare_counts_dataframe(raw_counts: Any) -> Any:
    pd = _import_dependency("pandas")

    counts = raw_counts.copy()
    if counts.empty:
        raise ValueError("Counts matrix is empty.")

    first_column = counts.columns[0]
    if (
        first_column not in counts.select_dtypes(include=["number"]).columns
        or str(first_column).lower() in {"gene", "gene_id", "genes", "symbol", "id", "entrezid"}
    ):
        counts = counts.set_index(first_column)

    counts.index = counts.index.map(str)
    counts.columns = [str(column).strip() for column in counts.columns]
    counts = counts.loc[:, ~pd.Index(counts.columns).duplicated()]
    counts = counts.apply(pd.to_numeric, errors="coerce")

    if counts.isna().all().all():
        raise ValueError("Counts matrix does not contain numeric sample columns.")

    counts = counts.fillna(0)
    counts = counts.groupby(level=0).sum()
    counts = counts.loc[counts.sum(axis=1) > 0]
    counts = counts.loc[:, counts.sum(axis=0) > 0]
    return counts


def _detect_sample_column(metadata: Any, sample_names: list[str]) -> str:
    candidates = [
        "sample_id",
        "sample",
        "sample_name",
        "Sample",
        "SampleID",
        "sampleid",
        "run",
    ]

    for candidate in candidates:
        if candidate in metadata.columns:
            return candidate

    best_column = ""
    best_overlap = -1
    sample_name_set = {str(name) for name in sample_names}
    for column in metadata.columns:
        overlap = sum(str(value) in sample_name_set for value in metadata[column].astype(str))
        if overlap > best_overlap:
            best_overlap = overlap
            best_column = str(column)

    if best_column and best_overlap > 0:
        return best_column

    raise ValueError("Could not detect sample identifier column in metadata.")


def _normalize_metadata(metadata: Any) -> Any:
    metadata = metadata.copy()
    metadata.columns = [str(column).strip() for column in metadata.columns]
    return metadata


def _compute_log_cpm(counts: Any) -> Any:
    np = _import_dependency("numpy")

    library_sizes = counts.sum(axis=0).replace(0, np.nan)
    counts_per_million = counts.divide(library_sizes, axis=1) * 1_000_000
    return np.log2(counts_per_million.fillna(0) + 1.0)


def _infer_gene_labels(gene_ids: list[str]) -> list[str]:
    mygene = None
    try:
        mygene = _import_dependency("mygene")
    except RuntimeError:
        return gene_ids

    preview_gene_ids = [str(gene_id).strip() for gene_id in gene_ids[:100] if str(gene_id).strip()]
    looks_like_ensembl = any(gene_id.upper().startswith("ENS") for gene_id in preview_gene_ids[:20])
    looks_like_entrez = bool(preview_gene_ids) and all(gene_id.isdigit() for gene_id in preview_gene_ids[:20])
    if not looks_like_ensembl and not looks_like_entrez:
        return gene_ids

    scopes = "ensembl.gene" if looks_like_ensembl else "entrezgene"

    try:
        mg = mygene.MyGeneInfo()
        query_results = mg.querymany(
            gene_ids,
            scopes=scopes,
            fields="symbol",
            species="all",
            as_dataframe=False,
            verbose=False,
        )
    except Exception as exc:  # pragma: no cover - external service failure
        LOGGER.warning("Could not annotate Ensembl identifiers with mygene: %s", exc)
        return gene_ids

    symbol_by_query = {
        str(item.get("query")): str(item.get("symbol"))
        for item in query_results
        if isinstance(item, dict) and item.get("query") and item.get("symbol")
    }
    return [symbol_by_query.get(gene_id, gene_id) for gene_id in gene_ids]


def _encode_image_data_uri(path: Path) -> str | None:
    if not path.exists():
        return None

    suffix = path.suffix.lower()
    mime_type = {
        ".png": "image/png",
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".webp": "image/webp",
        ".svg": "image/svg+xml",
    }.get(suffix)
    if not mime_type:
        return None

    return f"data:{mime_type};base64,{base64.b64encode(path.read_bytes()).decode('ascii')}"


def _resolve_report_branding(author: str | None, entity_name: str | None, entity_logo_path: str | None) -> ReportBranding:
    normalized_author = str(author or "").strip() or "Juan Vladimir de la Rosa Medina"
    normalized_entity_name = str(entity_name or "").strip()
    normalized_logo_path = str(entity_logo_path or "").strip()
    return ReportBranding(
        author=normalized_author,
        entity_name=normalized_entity_name,
        entity_logo_path=normalized_logo_path,
    )


def _resolve_atom_logo_path() -> Path | None:
    candidate = Path(__file__).resolve().parents[1] / "frontend" / "public" / "images" / "logo.png"
    if candidate.exists() and candidate.is_file():
        return candidate
    return None


def _save_summary_json(output_dir: Path, summary_payload: dict[str, Any]) -> Path:
    summary_path = output_dir / "summary_metrics.json"
    summary_path.write_text(json.dumps(summary_payload, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")
    return summary_path


def _safe_float(value: Any, digits: int = 3) -> str:
    try:
        numeric = float(value)
    except (TypeError, ValueError):
        return ""
    if math.isnan(numeric) or math.isinf(numeric):
        return ""
    return f"{numeric:.{digits}g}"


def _format_gene_label(row: Any) -> str:
    gene_symbol = str(row.get("gene_symbol") or "").strip()
    gene_id = str(row.get("gene_id") or "").strip()
    if gene_symbol and gene_symbol != gene_id:
        return f"{gene_symbol} ({gene_id})"
    return gene_symbol or gene_id


def _build_gene_summary(frame: Any, limit: int = 8) -> list[str]:
    if frame.empty:
        return []
    return [_format_gene_label(row) for _, row in frame.head(limit).iterrows()]


def _build_ranked_gene_details(frame: Any, limit: int = 8) -> list[str]:
    if frame.empty:
        return []
    ranked = frame.head(limit).copy()
    details = []
    for _, row in ranked.iterrows():
        details.append(
            ", ".join(
                part
                for part in (
                    _format_gene_label(row),
                    f"log2FC={_safe_float(row.get('log2FoldChange'))}",
                    f"padj={_safe_float(row.get('padj'))}",
                )
                if part
            )
        )
    return details


def _build_sample_distribution(metadata: Any, condition_column: str) -> list[dict[str, Any]]:
    distribution = metadata[condition_column].astype(str).value_counts(dropna=False)
    return [
        {"group": str(group), "count": int(count)}
        for group, count in distribution.items()
    ]


def _build_enrichment_highlights(enrichment_results: dict[str, Any], limit: int = 5) -> dict[str, list[dict[str, str]]]:
    highlights: dict[str, list[dict[str, str]]] = {}
    for key, frame in enrichment_results.items():
        if getattr(frame, "empty", True):
            continue

        prepared = frame.copy()
        description_column = next(
            (column for column in ("Term", "Description", "NAME") if column in prepared.columns),
            None,
        )
        score_column = next(
            (column for column in ("Adjusted P-value", "FDR q-val", "NOM p-val", "NES") if column in prepared.columns),
            None,
        )
        if description_column is None:
            continue

        if score_column and score_column in prepared.columns:
            prepared[score_column] = prepared[score_column]
        highlights[key] = [
            {
                "term": str(row.get(description_column) or "").strip(),
                "score": _safe_float(row.get(score_column)) if score_column else "",
            }
            for _, row in prepared.head(limit).iterrows()
            if str(row.get(description_column) or "").strip()
        ]
    return highlights


def _join_reference_lines(lines: list[str]) -> str:
    return " ".join([line for line in lines if line])


def _build_python_curated_references() -> list[str]:
    return [
        "Love MI, Huber W, Anders S. Moderated estimation of fold change and dispersion for RNA-seq data with DESeq2. Genome Biology. 2014;15(12):550.",
        "Subramanian A, Tamayo P, Mootha VK, et al. Gene set enrichment analysis: a knowledge-based approach for interpreting genome-wide expression profiles. Proceedings of the National Academy of Sciences USA. 2005;102(43):15545-15550.",
        "Kuleshov MV, Jones MR, Rouillard AD, et al. Enrichr: a comprehensive gene set enrichment analysis web server 2016 update. Nucleic Acids Research. 2016;44(W1):W90-W97.",
    ]


def _build_image_entries(generated_assets: dict[str, Path | None]) -> list[dict[str, str]]:
    entries: list[dict[str, str]] = []
    for key, path in generated_assets.items():
        if path is None or not path.exists():
            continue
        metadata = FIGURE_METADATA.get(
            key,
            {
                "label": key.replace("_", " ").title(),
                "caption": "Generated figure from Python RNA-seq workflow.",
                "group": "Additional figures",
            },
        )
        entries.append(
            {
                "filename": path.name,
                "label": metadata["label"],
                "caption": metadata["caption"],
                "group": metadata["group"],
                "src": _encode_image_data_uri(path) or path.name,
            }
        )
    return entries


def _build_report_sections() -> list[dict[str, str]]:
    return [
        {"id": "section-introduccion", "title": "0.1 Introducción"},
        {"id": "section-sample-distribution", "title": "0.2 Información de las muestras"},
        {"id": "section-quality-control", "title": "0.3 Control de calidad"},
        {"id": "section-de-overview", "title": "0.4 Expresión Génica Diferencial"},
        {"id": "section-top-hits", "title": "0.5 Genes destacados"},
        {"id": "section-figure-overview", "title": "0.6 Resumen de figuras"},
        {"id": "section-qc-figures", "title": "0.7 Figuras de control de calidad"},
        {"id": "section-de-figures", "title": "0.8 Figuras de expresión diferencial"},
        {"id": "section-enrichment-figures", "title": "0.9 Figuras de enriquecimiento funcional"},
        {"id": "section-functional-enrichment", "title": "0.10 Enriquecimiento funcional"},
        {"id": "section-interpretation", "title": "0.11 Interpretación biológica integrada"},
        {"id": "section-validation", "title": "0.12 Validación o siguientes pasos"},
        {"id": "section-limitations", "title": "0.13 Limitaciones"},
        {"id": "section-enrichment-summary", "title": "0.14 Resumen de enriquecimiento"},
        {"id": "section-references", "title": "0.15 Referencias"},
    ]


def _find_first_column(frame: Any, candidates: list[str]) -> str | None:
    for candidate in candidates:
        if candidate in frame.columns:
            return candidate
    return None


def _limit_prerank_frame(frame: Any, max_genes: int = 5000) -> Any:
    if getattr(frame, "empty", True) or len(frame) <= max_genes:
        return frame

    pd = _import_dependency("pandas")
    limited = frame.copy()
    limited["abs_score"] = pd.to_numeric(limited["score"], errors="coerce").abs()
    limited = limited.sort_values(by="abs_score", ascending=False, na_position="last").head(max_genes).copy()
    return limited.drop(columns=["abs_score"], errors="ignore")


def _configure_document_styles(document: Any, docx_module: Any) -> None:
    section = document.sections[0]
    section.top_margin = docx_module.shared.Inches(0.7)
    section.bottom_margin = docx_module.shared.Inches(0.7)
    section.left_margin = docx_module.shared.Inches(0.8)
    section.right_margin = docx_module.shared.Inches(0.8)

    styles = document.styles
    normal_style = styles["Normal"]
    normal_style.font.name = "Aptos"
    normal_style.font.size = docx_module.shared.Pt(10.5)
    normal_style.font.color.rgb = docx_module.shared.RGBColor(30, 41, 59)

    title_style = styles["Title"]
    title_style.font.name = "Aptos Display"
    title_style.font.size = docx_module.shared.Pt(24)
    title_style.font.bold = True
    title_style.font.color.rgb = docx_module.shared.RGBColor(15, 23, 42)
    title_style.paragraph_format.alignment = docx_module.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_after = docx_module.shared.Pt(8)

    subtitle_style = styles["Subtitle"]
    subtitle_style.font.name = "Aptos"
    subtitle_style.font.size = docx_module.shared.Pt(11)
    subtitle_style.font.color.rgb = docx_module.shared.RGBColor(71, 85, 105)
    subtitle_style.paragraph_format.alignment = docx_module.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    subtitle_style.paragraph_format.space_after = docx_module.shared.Pt(14)

    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        style = styles[style_name]
        style.font.name = "Aptos Display"
        style.font.color.rgb = docx_module.shared.RGBColor(15, 23, 42)

    styles["Heading 1"].font.size = docx_module.shared.Pt(16)
    styles["Heading 1"].font.bold = True
    styles["Heading 2"].font.size = docx_module.shared.Pt(13)
    styles["Heading 2"].font.bold = True
    styles["Heading 3"].font.size = docx_module.shared.Pt(11.5)
    styles["Heading 3"].font.bold = True


def _add_docx_section_heading(document: Any, text: str, level: int = 1) -> Any:
    heading = document.add_paragraph(style=f"Heading {level}")
    heading.paragraph_format.space_before = 0
    heading.paragraph_format.space_after = 0
    run = heading.add_run(text)
    run.bold = True
    return heading


def _add_docx_body_paragraph(document: Any, text: str, style: str | None = None) -> Any:
    paragraph = document.add_paragraph(style=style)
    paragraph.paragraph_format.space_after = 0
    paragraph.paragraph_format.line_spacing = 1.1
    paragraph.add_run(text)
    return paragraph


def _configure_docx_header(document: Any, output_dir: Path, branding: ReportBranding, docx_module: Any) -> None:
    atom_logo_path = _resolve_atom_logo_path()
    safe_atom_logo_path = _prepare_docx_image_path(atom_logo_path, output_dir) if atom_logo_path else None
    safe_entity_logo_path = (
        _prepare_docx_image_path(Path(branding.entity_logo_path), output_dir)
        if branding.entity_logo_path
        else None
    )

    for section in document.sections:
        header = section.header
        header.is_linked_to_previous = False
        paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        paragraph.alignment = docx_module.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        if safe_atom_logo_path and safe_atom_logo_path.exists():
            paragraph.add_run().add_picture(str(safe_atom_logo_path), width=docx_module.shared.Inches(0.55))
        if safe_entity_logo_path and safe_entity_logo_path.exists():
            paragraph.add_run("   ")
            paragraph.add_run().add_picture(str(safe_entity_logo_path), width=docx_module.shared.Inches(0.55))


def _prepare_docx_image_path(path: Path, output_dir: Path) -> Path | None:
    if not path.exists() or not path.is_file():
        return None

    if path.suffix.lower() in {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tif", ".tiff"}:
        return path

    try:
        image_module = _import_dependency("PIL.Image", "Pillow")
        with image_module.open(path) as image:
            converted_path = output_dir / f"{path.stem}_docx.png"
            if image.mode not in {"RGB", "RGBA"}:
                image = image.convert("RGBA")
            image.save(converted_path, format="PNG")
            return converted_path
    except Exception as exc:  # pragma: no cover - depends on installed codecs/assets
        LOGGER.warning("Could not convert image for DOCX %s: %s", path, exc)
        return None


def _add_docx_figure(document: Any, label: str, path: Path, docx_module: Any) -> None:
    metadata = FIGURE_METADATA.get(
        label,
        {
            "label": label.replace("_", " ").title(),
            "caption": "Generated figure from Python RNA-seq workflow.",
        },
    )
    safe_path = _prepare_docx_image_path(path, path.parent)
    if not safe_path:
        LOGGER.warning("Skipping DOCX figure because image could not be prepared: %s", path)
        return
    _add_docx_section_heading(document, metadata["label"], level=2)
    document.add_picture(str(safe_path), width=docx_module.shared.Inches(6.3))
    caption = document.add_paragraph(style="Subtitle")
    caption.paragraph_format.space_after = 0
    caption.alignment = docx_module.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    caption.add_run(metadata["caption"])


def _add_docx_metric_table(document: Any, summary: dict[str, Any], docx_module: Any) -> None:
    table = document.add_table(rows=3, cols=4)
    table.style = "Light List Accent 1"
    metrics = [
        ("Genes de entrada", summary.get("n_input_genes")),
        ("Muestras de entrada", summary.get("n_input_samples")),
        ("Genes evaluados", summary.get("n_tested_genes")),
        ("Genes significativos", summary.get("n_significant_genes")),
        ("Sobreexpresados", summary.get("n_upregulated")),
        ("Infraexpresados", summary.get("n_downregulated")),
    ]
    for index, (label, value) in enumerate(metrics):
        row = index // 2
        col = (index % 2) * 2
        table.cell(row, col).text = str(label)
        table.cell(row, col + 1).text = str(value)
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = 0
                for run in paragraph.runs:
                    run.font.size = docx_module.shared.Pt(10)
                    if cell.text in {str(label) for label, _ in metrics}:
                        run.bold = True


def _add_docx_divider(document: Any, docx_module: Any) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = docx_module.shared.Pt(6)
    paragraph.paragraph_format.space_after = docx_module.shared.Pt(8)
    run = paragraph.add_run(" ")
    border = paragraph._p.get_or_add_pPr()
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "D6DEE8")
    p_bdr.append(bottom)
    border.append(p_bdr)
    run.font.size = docx_module.shared.Pt(1)


def _resolve_figure_key_from_label(label: str) -> str:
    for key, metadata in FIGURE_METADATA.items():
        if metadata["label"] == label:
            return key
    return label.lower().replace(" ", "_")


def _format_term_score(row: dict[str, str]) -> str:
    suffix = f" (score={row['score']})" if row.get("score") else ""
    return f"{row['term']}{suffix}"


def _build_qc_summary(sample_distribution: list[dict[str, Any]]) -> str:
    if not sample_distribution:
        return "Sample distribution could not be summarised from metadata."
    groups = ", ".join(f"{item['group']} (n={item['count']})" for item in sample_distribution)
    return (
        f"Selected samples were balanced across groups: {groups}. "
        "PCA and heatmap figures should be interpreted together to confirm group separation and clustering consistency."
    )


def _build_functional_enrichment_summary(enrichment_highlights: dict[str, list[dict[str, str]]]) -> dict[str, str]:
    prerank_rows = enrichment_highlights.get("gsea_prerank", [])
    up_rows = enrichment_highlights.get("enrichr_up", [])
    down_rows = enrichment_highlights.get("enrichr_down", [])

    prerank_text = (
        "; ".join(
            _format_term_score(row)
            for row in prerank_rows[:4]
        )
        if prerank_rows
        else "No preranked terms passed current filters."
    )
    up_text = (
        "; ".join(
            _format_term_score(row)
            for row in up_rows[:4]
        )
        if up_rows
        else "No enriched terms were retained for upregulated genes."
    )
    down_text = (
        "; ".join(
            _format_term_score(row)
            for row in down_rows[:4]
        )
        if down_rows
        else "No enriched terms were retained for downregulated genes."
    )

    return {
        "prerank": prerank_text,
        "up": up_text,
        "down": down_text,
    }


def _build_integrated_interpretation(
    top_up_genes: list[str],
    top_down_genes: list[str],
    enrichment_summary: dict[str, str],
) -> str:
    up_text = ", ".join(top_up_genes[:5]) if top_up_genes else "upregulated candidates"
    down_text = ", ".join(top_down_genes[:5]) if top_down_genes else "downregulated candidates"
    interpretation_parts = [
        f"Upregulated genes such as {up_text} define a main axis of response in the evaluated contrast.",
        f"Downregulated genes such as {down_text} delimit a repressed biological programme across compared samples.",
    ]
    if "No preranked terms passed current filters." not in enrichment_summary["prerank"]:
        interpretation_parts.append(
            f"Pathway- and process-level signal in preranked analysis points to: {enrichment_summary['prerank']}."
        )
    if "No enriched terms were retained for upregulated genes." not in enrichment_summary["up"]:
        interpretation_parts.append(
            f"The upregulated programme is supported by gene sets such as {enrichment_summary['up']}."
        )
    if "No enriched terms were retained for downregulated genes." not in enrichment_summary["down"]:
        interpretation_parts.append(
            f"The downregulated programme is supported by gene sets such as {enrichment_summary['down']}."
        )
    return " ".join(interpretation_parts)


def _build_validation_suggestions(top_up_genes: list[str], top_down_genes: list[str]) -> list[str]:
    suggestions = []
    if top_up_genes:
        suggestions.append(
            f"Validate leading induced markers with qPCR or targeted protein assays, prioritising {', '.join(top_up_genes[:3])}."
        )
    if top_down_genes:
        suggestions.append(
            f"Confirm leading repressed markers with an orthogonal assay, prioritising {', '.join(top_down_genes[:3])}."
        )
    suggestions.append("Review quality-control figures together with metadata to confirm that no outlier is driving the observed separation.")
    suggestions.append("If biological interpretation is critical, repeat enrichment with alternative collections and a symbol-mapping audit.")
    return suggestions


def _build_limitations(summary: dict[str, Any], enrichment_highlights: dict[str, list[dict[str, str]]]) -> list[str]:
    limitations = [
        "Python report uses a complementary implementation and may not reproduce DESeq2 or clusterProfiler outputs exactly.",
        "Current implementation supports two-group contrasts and depends on correct identifier mapping for downstream enrichment.",
    ]
    if not enrichment_highlights:
        limitations.append("Functional enrichment remained limited under current filters and available libraries.")
    if int(summary.get("n_input_samples") or 0) < 8:
        limitations.append("Sample size is modest, so effect-size stability and dispersion estimates should be interpreted with caution.")
    return limitations


def load_inputs(counts_path: str, metadata_path: str) -> PipelineInputs:
    LOGGER.info("Loading counts matrix from %s", counts_path)
    LOGGER.info("Loading metadata from %s", metadata_path)

    counts_file = Path(counts_path).expanduser().resolve()
    metadata_file = Path(metadata_path).expanduser().resolve()

    if not counts_file.exists():
        raise FileNotFoundError(f"Counts file not found: {counts_file}")
    if not metadata_file.exists():
        raise FileNotFoundError(f"Metadata file not found: {metadata_file}")

    raw_counts = _read_tabular_file(counts_file)
    raw_metadata = _read_tabular_file(metadata_file)
    counts = _prepare_counts_dataframe(raw_counts)
    metadata = _normalize_metadata(raw_metadata)
    sample_column = _detect_sample_column(metadata, counts.columns.tolist())

    LOGGER.info("Loaded %s genes and %s samples from counts matrix", counts.shape[0], counts.shape[1])
    LOGGER.info("Detected metadata sample column: %s", sample_column)
    return PipelineInputs(counts=counts, metadata=metadata, sample_column=sample_column)


def validate_inputs(
    counts: Any,
    metadata: Any,
    condition_column: str,
    case: str,
    control: str,
    sample_column: str,
) -> tuple[Any, Any]:
    pd = _import_dependency("pandas")

    if condition_column not in metadata.columns:
        raise ValueError(f"Condition column '{condition_column}' not found in metadata.")

    selected_metadata = metadata[metadata[condition_column].astype(str).isin({case, control})].copy()
    if selected_metadata.empty:
        raise ValueError("No metadata rows match requested case/control contrast.")

    selected_metadata[sample_column] = selected_metadata[sample_column].astype(str)
    selected_metadata = selected_metadata.drop_duplicates(subset=[sample_column], keep="first")

    sample_ids = selected_metadata[sample_column].tolist()
    missing_in_counts = sorted(set(sample_ids) - set(counts.columns))
    if missing_in_counts:
        raise ValueError(f"Samples present in metadata but missing in counts: {missing_in_counts}")

    overlapping_samples = [sample_id for sample_id in counts.columns if sample_id in sample_ids]
    if len(overlapping_samples) < 2:
        raise ValueError("Need at least two overlapping samples between counts and metadata.")

    filtered_counts = counts.loc[:, overlapping_samples].copy()
    filtered_metadata = selected_metadata.set_index(sample_column).loc[overlapping_samples].copy()
    filtered_metadata[condition_column] = pd.Categorical(
        filtered_metadata[condition_column].astype(str),
        categories=[control, case],
        ordered=True,
    )

    group_sizes = filtered_metadata[condition_column].value_counts(dropna=False).to_dict()
    if group_sizes.get(control, 0) == 0 or group_sizes.get(case, 0) == 0:
        raise ValueError("Both case and control groups must contain at least one sample.")

    if min(group_sizes.get(control, 0), group_sizes.get(case, 0)) < 2:
        LOGGER.warning("Contrast has fewer than two replicates in one group. PyDESeq2 may be unstable.")

    LOGGER.info("Validated %s genes across %s selected samples", filtered_counts.shape[0], filtered_counts.shape[1])
    LOGGER.info("Group sizes: %s", group_sizes)
    return filtered_counts, filtered_metadata


def run_differential_expression(
    counts: Any,
    metadata: Any,
    condition_column: str,
    case: str,
    control: str,
    alpha: float = 0.05,
    log2fc_threshold: float = 1.0,
) -> DifferentialExpressionOutputs:
    np = _import_dependency("numpy")
    pd = _import_dependency("pandas")
    pydeseq2_dds = _import_dependency("pydeseq2.dds", "pydeseq2")
    pydeseq2_ds = _import_dependency("pydeseq2.ds", "pydeseq2")
    pydeseq2_default_inference = _import_dependency("pydeseq2.default_inference", "pydeseq2")

    LOGGER.info("Running differential expression with PyDESeq2")
    counts_for_deseq = counts.round().astype(int).T.copy()
    inference = pydeseq2_default_inference.DefaultInference(n_cpus=1, backend="loky")
    dds = pydeseq2_dds.DeseqDataSet(
        counts=counts_for_deseq,
        metadata=metadata,
        design_factors=condition_column,
        refit_cooks=True,
        n_cpus=1,
        inference=inference,
        low_memory=True,
    )
    dds.deseq2()

    stats = pydeseq2_ds.DeseqStats(
        dds,
        contrast=[condition_column, case, control],
        alpha=alpha,
        inference=inference,
        n_cpus=1,
    )
    stats.summary()
    results = stats.results_df.copy().reset_index()
    if "index" in results.columns:
        results = results.rename(columns={"index": "gene_id"})
    elif results.columns[0] != "gene_id":
        results = results.rename(columns={results.columns[0]: "gene_id"})

    results["gene_id"] = results["gene_id"].astype(str)
    results["gene_symbol"] = _infer_gene_labels(results["gene_id"].tolist())
    for column in ("log2FoldChange", "pvalue", "padj"):
        if column not in results.columns:
            raise ValueError(f"PyDESeq2 output missing expected column: {column}")
        results[column] = pd.to_numeric(results[column], errors="coerce")

    results = results.sort_values(by=["padj", "pvalue"], na_position="last").reset_index(drop=True)
    significance_mask = (
        results["padj"].notna()
        & (results["padj"] < alpha)
        & (results["log2FoldChange"].abs() >= log2fc_threshold)
    )
    significant = results.loc[significance_mask].copy()
    upregulated = significant.loc[significant["log2FoldChange"] >= log2fc_threshold].copy()
    downregulated = significant.loc[significant["log2FoldChange"] <= -log2fc_threshold].copy()

    ranked_scores = (
        results[["gene_id", "log2FoldChange"]]
        .dropna()
        .drop_duplicates(subset=["gene_id"], keep="first")
        .set_index("gene_id")["log2FoldChange"]
        .sort_values(ascending=False)
    )

    normalized_counts = _compute_log_cpm(counts)
    top_up = _build_gene_summary(upregulated.nsmallest(10, "padj"), limit=10)
    top_down = _build_gene_summary(downregulated.nsmallest(10, "padj"), limit=10)

    summary = {
        "n_input_genes": int(counts.shape[0]),
        "n_input_samples": int(counts.shape[1]),
        "n_tested_genes": int(results["pvalue"].notna().sum()),
        "n_significant_genes": int(significant.shape[0]),
        "n_upregulated": int(upregulated.shape[0]),
        "n_downregulated": int(downregulated.shape[0]),
        "top_upregulated_genes": top_up,
        "top_downregulated_genes": top_down,
        "alpha": alpha,
        "log2fc_threshold": log2fc_threshold,
        "min_log2fc": float(results["log2FoldChange"].min(skipna=True)) if not results.empty else math.nan,
        "max_log2fc": float(results["log2FoldChange"].max(skipna=True)) if not results.empty else math.nan,
    }

    LOGGER.info(
        "Differential expression completed: %s significant genes (%s up, %s down)",
        summary["n_significant_genes"],
        summary["n_upregulated"],
        summary["n_downregulated"],
    )
    return DifferentialExpressionOutputs(
        counts=counts,
        metadata=metadata,
        normalized_counts=normalized_counts,
        results=results,
        significant=significant,
        upregulated=upregulated,
        downregulated=downregulated,
        ranked_scores=ranked_scores,
        summary=summary,
    )


def generate_volcano_plot(
    results: Any,
    output_dir: Path,
    alpha: float = 0.05,
    log2fc_threshold: float = 1.0,
) -> Path:
    np = _import_dependency("numpy")
    matplotlib = _import_dependency("matplotlib")
    matplotlib.use("Agg")
    plt = _import_dependency("matplotlib.pyplot", "matplotlib")

    plot_frame = results.copy()
    plot_frame["minus_log10_padj"] = -np.log10(plot_frame["padj"].clip(lower=1e-300))
    plot_frame["status"] = "Not significant"
    plot_frame.loc[
        (plot_frame["padj"] < alpha) & (plot_frame["log2FoldChange"] >= log2fc_threshold),
        "status",
    ] = "Upregulated"
    plot_frame.loc[
        (plot_frame["padj"] < alpha) & (plot_frame["log2FoldChange"] <= -log2fc_threshold),
        "status",
    ] = "Downregulated"

    color_map = {
        "Not significant": "#94a3b8",
        "Upregulated": "#dc2626",
        "Downregulated": "#2563eb",
    }

    figure, axis = plt.subplots(figsize=(10, 8))
    for status, color in color_map.items():
        subset = plot_frame[plot_frame["status"] == status]
        axis.scatter(
            subset["log2FoldChange"],
            subset["minus_log10_padj"],
            s=18 if status == "Not significant" else 28,
            alpha=0.7,
            c=color,
            label=status,
            edgecolors="none",
        )

    axis.axvline(log2fc_threshold, linestyle="--", linewidth=1, color="#475569")
    axis.axvline(-log2fc_threshold, linestyle="--", linewidth=1, color="#475569")
    axis.axhline(-math.log10(alpha), linestyle="--", linewidth=1, color="#475569")
    axis.set_title("Volcano plot")
    axis.set_xlabel("log2 fold change")
    axis.set_ylabel("-log10 adjusted p-value")
    axis.legend(frameon=False)
    axis.grid(alpha=0.15)

    output_path = output_dir / "volcano_plot.png"
    figure.tight_layout()
    figure.savefig(output_path, dpi=220)
    plt.close(figure)
    LOGGER.info("Saved volcano plot to %s", output_path)
    return output_path


def generate_heatmap(
    normalized_counts: Any,
    significant_results: Any,
    metadata: Any,
    condition_column: str,
    output_dir: Path,
    top_n: int = 50,
) -> Path | None:
    if significant_results.empty:
        LOGGER.warning("Skipping heatmap because no significant genes were found.")
        return None

    np = _import_dependency("numpy")
    pd = _import_dependency("pandas")
    matplotlib = _import_dependency("matplotlib")
    matplotlib.use("Agg")
    plt = _import_dependency("matplotlib.pyplot", "matplotlib")
    sns = _import_dependency("seaborn")

    selected = (
        significant_results.sort_values(by=["padj", "pvalue"], na_position="last")
        .head(top_n)["gene_id"]
        .tolist()
    )
    heatmap_data = normalized_counts.loc[normalized_counts.index.intersection(selected)].copy()
    if heatmap_data.empty:
        LOGGER.warning("Skipping heatmap because selected genes were not present in normalized matrix.")
        return None

    heatmap_data = heatmap_data.subtract(heatmap_data.mean(axis=1), axis=0)
    heatmap_data = heatmap_data.divide(heatmap_data.std(axis=1).replace(0, np.nan), axis=0).fillna(0)
    row_labels = _infer_gene_labels(heatmap_data.index.tolist())
    heatmap_data.index = pd.Index(row_labels, name="gene_label")

    palette = {
        str(level): color
        for level, color in zip(
            metadata[condition_column].astype(str).unique().tolist(),
            ["#2563eb", "#dc2626", "#16a34a", "#d97706", "#7c3aed"],
            strict=False,
        )
    }
    column_colors = metadata[condition_column].astype(str).map(lambda value: palette.get(value, "#64748b"))

    cluster = sns.clustermap(
        heatmap_data,
        cmap="vlag",
        col_colors=column_colors,
        figsize=(12, 12),
        xticklabels=True,
        yticklabels=True,
        linewidths=0,
    )
    cluster.ax_heatmap.set_title("Heatmap of significant genes")

    output_path = output_dir / "significant_genes_heatmap.png"
    cluster.figure.savefig(output_path, dpi=220, bbox_inches="tight")
    plt.close(cluster.figure)
    LOGGER.info("Saved heatmap to %s", output_path)
    return output_path


def generate_pca(normalized_counts: Any, metadata: Any, condition_column: str, output_dir: Path) -> Path:
    pd = _import_dependency("pandas")
    matplotlib = _import_dependency("matplotlib")
    matplotlib.use("Agg")
    plt = _import_dependency("matplotlib.pyplot", "matplotlib")
    sns = _import_dependency("seaborn")
    decomposition = _import_dependency("sklearn.decomposition")

    pca = decomposition.PCA(n_components=2)
    transformed = pca.fit_transform(normalized_counts.T)
    pca_frame = pd.DataFrame(
        transformed,
        index=normalized_counts.columns,
        columns=["PC1", "PC2"],
    ).join(metadata[[condition_column]])

    figure, axis = plt.subplots(figsize=(9, 7))
    sns.scatterplot(
        data=pca_frame,
        x="PC1",
        y="PC2",
        hue=condition_column,
        s=90,
        palette="deep",
        ax=axis,
    )
    axis.set_title("PCA of samples")
    axis.set_xlabel(f"PC1 ({pca.explained_variance_ratio_[0] * 100:.1f}% variance)")
    axis.set_ylabel(f"PC2 ({pca.explained_variance_ratio_[1] * 100:.1f}% variance)")
    axis.grid(alpha=0.15)
    axis.legend(frameon=False)

    output_path = output_dir / "pca_samples.png"
    figure.tight_layout()
    figure.savefig(output_path, dpi=220)
    plt.close(figure)
    LOGGER.info("Saved PCA plot to %s", output_path)
    return output_path


def generate_sample_distance_heatmap(normalized_counts: Any, metadata: Any, condition_column: str, output_dir: Path) -> Path:
    np = _import_dependency("numpy")
    pd = _import_dependency("pandas")
    matplotlib = _import_dependency("matplotlib")
    matplotlib.use("Agg")
    plt = _import_dependency("matplotlib.pyplot", "matplotlib")
    sns = _import_dependency("seaborn")

    sample_matrix = normalized_counts.T
    correlation_matrix = sample_matrix.T.corr(method="pearson")
    distance_matrix = 1 - correlation_matrix
    distance_matrix.index = pd.Index(sample_matrix.index, name="sample")
    distance_matrix.columns = pd.Index(sample_matrix.index, name="sample")

    palette = {
        str(level): color
        for level, color in zip(
            metadata[condition_column].astype(str).unique().tolist(),
            ["#2563eb", "#dc2626", "#16a34a", "#d97706", "#7c3aed"],
            strict=False,
        )
    }
    sample_groups = metadata.loc[sample_matrix.index, condition_column].astype(str)
    sample_colors = sample_groups.map(lambda value: palette.get(value, "#64748b"))

    cluster = sns.clustermap(
        distance_matrix,
        cmap="mako",
        row_colors=sample_colors,
        col_colors=sample_colors,
        figsize=(10, 10),
        linewidths=0.2,
        xticklabels=True,
        yticklabels=True,
    )
    cluster.ax_heatmap.set_title("Sample distance heatmap")

    output_path = output_dir / "sample_distance_heatmap.png"
    cluster.figure.savefig(output_path, dpi=220, bbox_inches="tight")
    plt.close(cluster.figure)
    LOGGER.info("Saved sample distance heatmap to %s", output_path)
    return output_path


def generate_library_size_barplot(counts: Any, metadata: Any, condition_column: str, output_dir: Path) -> Path:
    pd = _import_dependency("pandas")
    matplotlib = _import_dependency("matplotlib")
    matplotlib.use("Agg")
    plt = _import_dependency("matplotlib.pyplot", "matplotlib")
    sns = _import_dependency("seaborn")

    library_sizes = counts.sum(axis=0)
    plot_frame = pd.DataFrame(
        {
            "sample": library_sizes.index.astype(str),
            "library_size": library_sizes.values,
        }
    ).set_index("sample")
    plot_frame = plot_frame.join(metadata[[condition_column]], how="left").reset_index()

    figure, axis = plt.subplots(figsize=(10, 6))
    sns.barplot(
        data=plot_frame,
        x="sample",
        y="library_size",
        hue=condition_column,
        dodge=False,
        palette="deep",
        ax=axis,
    )
    axis.set_title("Library size by sample")
    axis.set_xlabel("Sample")
    axis.set_ylabel("Raw counts")
    axis.tick_params(axis="x", rotation=25)
    axis.grid(axis="y", alpha=0.15)
    axis.legend(frameon=False)

    output_path = output_dir / "library_size_barplot.png"
    figure.tight_layout()
    figure.savefig(output_path, dpi=220, bbox_inches="tight")
    plt.close(figure)
    LOGGER.info("Saved library size bar plot to %s", output_path)
    return output_path


def generate_normalized_expression_boxplot(normalized_counts: Any, metadata: Any, condition_column: str, output_dir: Path) -> Path:
    pd = _import_dependency("pandas")
    matplotlib = _import_dependency("matplotlib")
    matplotlib.use("Agg")
    plt = _import_dependency("matplotlib.pyplot", "matplotlib")
    sns = _import_dependency("seaborn")

    log_counts = (normalized_counts + 1).applymap(lambda value: math.log2(float(value)))
    plot_frame = (
        log_counts.T
        .reset_index(names="sample")
        .melt(id_vars="sample", var_name="gene_id", value_name="log2_expression")
    )
    metadata_frame = metadata[[condition_column]].reset_index(names="sample")
    plot_frame = plot_frame.merge(metadata_frame, on="sample", how="left")

    figure, axis = plt.subplots(figsize=(11, 6))
    sns.boxplot(
        data=plot_frame,
        x="sample",
        y="log2_expression",
        hue=condition_column,
        dodge=False,
        palette="deep",
        linewidth=1,
        fliersize=1,
        ax=axis,
    )
    axis.set_title("Normalized expression distribution by sample")
    axis.set_xlabel("Sample")
    axis.set_ylabel("log2 normalized counts")
    axis.tick_params(axis="x", rotation=25)
    axis.grid(axis="y", alpha=0.15)
    axis.legend(frameon=False)

    output_path = output_dir / "normalized_expression_boxplot.png"
    figure.tight_layout()
    figure.savefig(output_path, dpi=220, bbox_inches="tight")
    plt.close(figure)
    LOGGER.info("Saved normalized expression box plot to %s", output_path)
    return output_path


def generate_ma_plot(results: Any, output_dir: Path, alpha: float = 0.05) -> Path | None:
    if "baseMean" not in results.columns:
        LOGGER.warning("Skipping MA plot because baseMean column is unavailable.")
        return None

    np = _import_dependency("numpy")
    matplotlib = _import_dependency("matplotlib")
    matplotlib.use("Agg")
    plt = _import_dependency("matplotlib.pyplot", "matplotlib")

    plot_frame = results.copy()
    plot_frame["mean_expression"] = np.log10(plot_frame["baseMean"].clip(lower=1) + 1)
    plot_frame["status"] = "Not significant"
    plot_frame.loc[
        (plot_frame["padj"] < alpha) & (plot_frame["log2FoldChange"] > 0),
        "status",
    ] = "Upregulated"
    plot_frame.loc[
        (plot_frame["padj"] < alpha) & (plot_frame["log2FoldChange"] < 0),
        "status",
    ] = "Downregulated"

    color_map = {
        "Not significant": "#94a3b8",
        "Upregulated": "#dc2626",
        "Downregulated": "#2563eb",
    }

    figure, axis = plt.subplots(figsize=(10, 7))
    for status, color in color_map.items():
        subset = plot_frame[plot_frame["status"] == status]
        axis.scatter(
            subset["mean_expression"],
            subset["log2FoldChange"],
            s=18 if status == "Not significant" else 24,
            alpha=0.7,
            c=color,
            label=status,
            edgecolors="none",
        )

    axis.axhline(0, linestyle="--", linewidth=1, color="#475569")
    axis.set_title("MA plot")
    axis.set_xlabel("log10 mean normalized expression")
    axis.set_ylabel("log2 fold change")
    axis.legend(frameon=False)
    axis.grid(alpha=0.15)

    output_path = output_dir / "ma_plot.png"
    figure.tight_layout()
    figure.savefig(output_path, dpi=220)
    plt.close(figure)
    LOGGER.info("Saved MA plot to %s", output_path)
    return output_path


def generate_top_genes_barplot(significant_results: Any, output_dir: Path, top_n_per_group: int = 6) -> Path | None:
    if significant_results.empty:
        LOGGER.warning("Skipping top significant genes bar plot because no significant genes were found.")
        return None

    pd = _import_dependency("pandas")
    matplotlib = _import_dependency("matplotlib")
    matplotlib.use("Agg")
    plt = _import_dependency("matplotlib.pyplot", "matplotlib")
    sns = _import_dependency("seaborn")

    ranked = significant_results.sort_values(by=["padj", "pvalue"], na_position="last").copy()
    up = ranked[ranked["log2FoldChange"] > 0].head(top_n_per_group).copy()
    down = ranked[ranked["log2FoldChange"] < 0].head(top_n_per_group).copy()
    selected = pd.concat([up, down], axis=0)
    if selected.empty:
        LOGGER.warning("Skipping top significant genes bar plot because no directional genes were found.")
        return None

    selected["gene_label"] = [
        _format_gene_label(row)
        for _, row in selected.iterrows()
    ]
    selected["direction"] = selected["log2FoldChange"].apply(lambda value: "Upregulated" if value > 0 else "Downregulated")
    selected = selected.sort_values(by="log2FoldChange")

    figure, axis = plt.subplots(figsize=(11, 8))
    sns.barplot(
        data=selected,
        x="log2FoldChange",
        y="gene_label",
        hue="direction",
        dodge=False,
        palette={"Upregulated": "#dc2626", "Downregulated": "#2563eb"},
        ax=axis,
    )
    axis.set_title("Top significant genes")
    axis.set_xlabel("log2 fold change")
    axis.set_ylabel("Gene")
    axis.legend(frameon=False)
    axis.grid(axis="x", alpha=0.15)

    output_path = output_dir / "top_significant_genes_barplot.png"
    figure.tight_layout()
    figure.savefig(output_path, dpi=220, bbox_inches="tight")
    plt.close(figure)
    LOGGER.info("Saved top significant genes bar plot to %s", output_path)
    return output_path


def generate_pvalue_histogram(results: Any, output_dir: Path) -> Path:
    matplotlib = _import_dependency("matplotlib")
    matplotlib.use("Agg")
    plt = _import_dependency("matplotlib.pyplot", "matplotlib")

    plot_series = results["padj"].dropna()
    figure, axis = plt.subplots(figsize=(9, 6))
    axis.hist(plot_series, bins=30, color="#2563eb", alpha=0.8, edgecolor="white")
    axis.set_title("Adjusted p-value distribution")
    axis.set_xlabel("Adjusted p-value")
    axis.set_ylabel("Gene count")
    axis.grid(axis="y", alpha=0.15)

    output_path = output_dir / "pvalue_histogram.png"
    figure.tight_layout()
    figure.savefig(output_path, dpi=220, bbox_inches="tight")
    plt.close(figure)
    LOGGER.info("Saved adjusted p-value histogram to %s", output_path)
    return output_path


def generate_enrichment_dotplot(
    frame: Any,
    output_dir: Path,
    output_name: str,
    title: str,
    top_n: int = 10,
) -> Path | None:
    if getattr(frame, "empty", True):
        LOGGER.warning("Skipping %s because enrichment frame is empty.", title)
        return None

    np = _import_dependency("numpy")
    pd = _import_dependency("pandas")
    matplotlib = _import_dependency("matplotlib")
    matplotlib.use("Agg")
    plt = _import_dependency("matplotlib.pyplot", "matplotlib")
    sns = _import_dependency("seaborn")

    term_column = _find_first_column(frame, ["Term", "Description", "NAME"])
    score_column = _find_first_column(frame, ["Adjusted P-value", "FDR q-val", "qvalue", "NOM p-val", "P-value"])
    effect_column = _find_first_column(frame, ["NES", "Combined Score", "Odds Ratio"])
    genes_column = _find_first_column(frame, ["Lead_genes", "Genes"])

    if term_column is None:
        LOGGER.warning("Skipping %s because no term/description column was found.", title)
        return None

    plot_frame = frame.copy()
    if score_column is not None:
        plot_frame = plot_frame.sort_values(by=score_column, ascending=True, na_position="last")
    plot_frame = plot_frame.head(top_n).copy()
    if plot_frame.empty:
        LOGGER.warning("Skipping %s because selected rows are empty.", title)
        return None

    plot_frame["term_label"] = plot_frame[term_column].astype(str).str.replace(r"\s+\(GO:\d+\)$", "", regex=True)
    plot_frame["term_label"] = plot_frame["term_label"].str.slice(0, 70)

    if score_column is not None:
        numeric_score = pd.to_numeric(plot_frame[score_column], errors="coerce").fillna(1.0)
        plot_frame["plot_score"] = -np.log10(numeric_score.clip(lower=1e-300))
    else:
        plot_frame["plot_score"] = np.arange(len(plot_frame), 0, -1)

    if effect_column is not None:
        numeric_effect = pd.to_numeric(plot_frame[effect_column], errors="coerce").fillna(0.0)
        plot_frame["effect_value"] = numeric_effect
        plot_frame["bubble_size"] = numeric_effect.abs().clip(lower=0.1) * 80
    else:
        plot_frame["effect_value"] = plot_frame["plot_score"]
        plot_frame["bubble_size"] = 120

    if genes_column is not None:
        plot_frame["gene_count"] = (
            plot_frame[genes_column]
            .astype(str)
            .apply(lambda value: len([item for item in re.split(r"[;,/]", value) if item.strip()]))
        )
        plot_frame["bubble_size"] = plot_frame["gene_count"].clip(lower=1) * 28

    plot_frame = plot_frame.sort_values(by="plot_score", ascending=True)

    figure_height = max(6, 0.55 * len(plot_frame) + 2.5)
    figure, axis = plt.subplots(figsize=(11, figure_height))
    scatter = axis.scatter(
        plot_frame["plot_score"],
        plot_frame["term_label"],
        s=plot_frame["bubble_size"],
        c=plot_frame["effect_value"],
        cmap="viridis",
        alpha=0.85,
        edgecolors="white",
        linewidths=0.8,
    )
    axis.set_title(title)
    axis.set_xlabel("-log10 significance" if score_column is not None else "rank")
    axis.set_ylabel("Term")
    axis.grid(axis="x", alpha=0.15)
    colorbar = figure.colorbar(scatter, ax=axis, pad=0.02)
    colorbar.set_label(effect_column or "effect")

    output_path = output_dir / output_name
    figure.tight_layout()
    figure.savefig(output_path, dpi=220, bbox_inches="tight")
    plt.close(figure)
    LOGGER.info("Saved enrichment dot plot to %s", output_path)
    return output_path


def run_enrichment(
    de_outputs: DifferentialExpressionOutputs,
    output_dir: Path,
) -> dict[str, Any]:
    pd = _import_dependency("pandas")
    enrichment_results: dict[str, Any] = {}

    try:
        gseapy = _import_dependency("gseapy")
    except RuntimeError as exc:
        LOGGER.warning("Skipping enrichment because gseapy is unavailable: %s", exc)
        return enrichment_results

    ranked_gene_ids = de_outputs.ranked_scores.index.tolist()
    ranked_gene_labels = _infer_gene_labels(ranked_gene_ids)
    prerank_frame = pd.DataFrame(
        {
            "gene_name": [str(label).strip().upper() for label in ranked_gene_labels],
            "score": de_outputs.ranked_scores.values,
        }
    )
    prerank_frame = prerank_frame.loc[prerank_frame["gene_name"].astype(bool)]
    prerank_frame = prerank_frame.drop_duplicates(subset=["gene_name"], keep="first")
    prerank_frame = _limit_prerank_frame(prerank_frame, max_genes=5000)

    gene_sets = [
        "GO_Biological_Process_2023",
        "GO_Molecular_Function_2023",
        "GO_Cellular_Component_2023",
    ]

    try:
        prerank_result = gseapy.prerank(
            rnk=prerank_frame,
            gene_sets=["GO_Biological_Process_2023"],
            outdir=None,
            permutation_num=40,
            min_size=5,
            max_size=500,
            threads=1,
            graph_num=0,
            no_plot=True,
            seed=42,
            verbose=False,
        )
        enrichment_results["gsea_prerank"] = prerank_result.res2d.copy()
        LOGGER.info("Completed prerank enrichment with %s rows", enrichment_results["gsea_prerank"].shape[0])
    except Exception as exc:  # pragma: no cover - depends on remote gene sets/runtime
        LOGGER.warning("Prerank enrichment failed: %s", exc)

    for direction, frame in {
        "up": de_outputs.upregulated,
        "down": de_outputs.downregulated,
    }.items():
        if frame.empty:
            continue

        gene_names = [str(label).strip().upper() for label in _infer_gene_labels(frame["gene_id"].astype(str).tolist())]
        gene_names = [label for label in gene_names if label]
        if len(gene_names) < 5:
            LOGGER.warning("Skipping %s enrichment because fewer than 5 genes are available.", direction)
            continue

        try:
            enrichr_result = gseapy.enrichr(
                gene_list=gene_names,
                gene_sets=gene_sets,
                outdir=None,
                cutoff=0.5,
                no_plot=True,
            )
            enrichment_results[f"enrichr_{direction}"] = enrichr_result.results.copy()
            LOGGER.info(
                "Completed %s enrichr enrichment with %s rows",
                direction,
                enrichment_results[f"enrichr_{direction}"].shape[0],
            )
        except Exception as exc:  # pragma: no cover - depends on remote gene sets/runtime
            LOGGER.warning("Enrichr %s enrichment failed: %s", direction, exc)

    return enrichment_results


def export_results(
    de_outputs: DifferentialExpressionOutputs,
    enrichment_results: dict[str, Any],
    output_dir: Path,
) -> dict[str, Path]:
    pd = _import_dependency("pandas")

    output_dir.mkdir(parents=True, exist_ok=True)
    exported_paths: dict[str, Path] = {}

    differential_path = output_dir / "differential_expression_results.tsv"
    de_outputs.results.to_csv(differential_path, sep="\t", index=False)
    exported_paths["differential_expression_tsv"] = differential_path

    significant_path = output_dir / "significant_genes.tsv"
    de_outputs.significant.to_csv(significant_path, sep="\t", index=False)
    exported_paths["significant_genes_tsv"] = significant_path

    summary_path = _save_summary_json(output_dir, de_outputs.summary)
    exported_paths["summary_json"] = summary_path

    workbook_path = output_dir / "rna_seq_python_results.xlsx"
    with pd.ExcelWriter(workbook_path, engine="openpyxl") as writer:
        de_outputs.results.to_excel(writer, sheet_name="DE_results", index=False)
        de_outputs.significant.to_excel(writer, sheet_name="Significant", index=False)
        de_outputs.upregulated.to_excel(writer, sheet_name="Upregulated", index=False)
        de_outputs.downregulated.to_excel(writer, sheet_name="Downregulated", index=False)
        de_outputs.metadata.reset_index().to_excel(writer, sheet_name="Metadata_used", index=False)
        for sheet_name, frame in enrichment_results.items():
            safe_name = sheet_name[:31]
            frame.to_excel(writer, sheet_name=safe_name, index=False)
    exported_paths["xlsx"] = workbook_path

    LOGGER.info("Exported tabular results to %s", output_dir)
    return exported_paths


def generate_html_report(
    output_dir: Path,
    de_outputs: DifferentialExpressionOutputs,
    generated_assets: dict[str, Path | None],
    enrichment_results: dict[str, Any],
    condition_column: str,
    case: str,
    control: str,
    branding: ReportBranding,
) -> Path:
    template_module = _import_dependency("jinja2")

    top_hits = de_outputs.significant.head(10).copy()
    top_hits_records = top_hits.fillna("").to_dict(orient="records")
    summary = de_outputs.summary
    sample_distribution = _build_sample_distribution(de_outputs.metadata, condition_column)
    image_entries = _build_image_entries(generated_assets)
    qc_images = [entry for entry in image_entries if entry["group"] == "Control de calidad"]
    de_images = [entry for entry in image_entries if entry["group"] == "Expresión diferencial"]
    enrichment_images = [entry for entry in image_entries if entry["group"] == "Enriquecimiento funcional"]
    other_images = [entry for entry in image_entries if entry["group"] not in {"Control de calidad", "Expresión diferencial", "Enriquecimiento funcional"}]
    enrichment_overview = {
        key: int(frame.shape[0])
        for key, frame in enrichment_results.items()
    }
    enrichment_highlights = _build_enrichment_highlights(enrichment_results)
    enrichment_summary = _build_functional_enrichment_summary(enrichment_highlights)
    top_up_genes = [str(gene) for gene in summary.get("top_upregulated_genes", [])[:10]]
    top_down_genes = [str(gene) for gene in summary.get("top_downregulated_genes", [])[:10]]
    top_up_details = _build_ranked_gene_details(de_outputs.upregulated.nsmallest(6, "padj"), limit=6)
    top_down_details = _build_ranked_gene_details(de_outputs.downregulated.nsmallest(6, "padj"), limit=6)
    qc_summary = _build_qc_summary(sample_distribution)
    integrated_interpretation = _build_integrated_interpretation(
        top_up_genes,
        top_down_genes,
        enrichment_summary,
    )
    validation_suggestions = _build_validation_suggestions(top_up_genes, top_down_genes)
    limitations = _build_limitations(summary, enrichment_highlights)
    curated_references = _build_python_curated_references()
    report_sections = _build_report_sections()
    report_design_id = output_dir.name.split("__")[0] if "__" in output_dir.name else output_dir.name
    generated_at = datetime.now().strftime("%Y-%m-%d %H:%M")
    report_author = branding.author
    report_entity_name = branding.entity_name
    atom_logo_src = ""
    atom_logo_path = _resolve_atom_logo_path()
    if atom_logo_path:
        atom_logo_src = _encode_image_data_uri(atom_logo_path) or ""
    entity_logo_src = ""
    if branding.entity_logo_path:
        entity_logo_src = _encode_image_data_uri(Path(branding.entity_logo_path)) or ""
    intro_text = (
        f"Este informe resume análisis de RNA-seq para contraste {case} frente a {control} "
        f"usando columna experimental {condition_column}. Se procesaron {summary['n_input_samples']} muestras "
        f"y {summary['n_tested_genes']} genes evaluables para caracterizar separación entre grupos, "
        "expresión diferencial y señales de enriquecimiento funcional comparables con flujo en R."
    )

    template = template_module.Template(
        """
<!doctype html>
<html lang="en">
  <head>
    <meta charset="utf-8" />
    <title>Análisis de RNA-seq</title>
    <style>
      body { font-family: Arial, sans-serif; margin: 0; padding: 2rem; background: linear-gradient(180deg, #e2e8f0 0%, #f8fafc 100%); color: #0f172a; line-height: 1.6; box-sizing: border-box; }
      *, *::before, *::after { box-sizing: inherit; }
      .page-shell { max-width: 1360px; margin: 0 auto; }
      .layout { display: grid; grid-template-columns: minmax(260px, 290px) minmax(0, 1fr); gap: 1.75rem; align-items: start; }
      .sidebar { position: sticky; top: 24px; align-self: start; border: 1px solid #cbd5e1; border-radius: 18px; background: linear-gradient(180deg, #ffffff 0%, #f8fbff 100%); box-shadow: 0 18px 40px rgba(15, 23, 42, 0.08); overflow: hidden; }
      .sidebar-hero { padding: 1.2rem 1.1rem 1rem; border-bottom: 1px solid #e2e8f0; background: linear-gradient(135deg, #0f172a 0%, #0f4c81 100%); color: #fff; }
      .sidebar-hero p { margin: 0; }
      .sidebar-logo-row { display: flex; gap: 0.6rem; align-items: center; flex-wrap: wrap; margin-bottom: 0.8rem; }
      .sidebar-logo { width: 62px; height: 62px; object-fit: contain; display: block; background: rgba(255,255,255,0.96); border-radius: 14px; padding: 0.35rem; }
      .sidebar-eyebrow { font-size: 0.76rem; text-transform: uppercase; letter-spacing: 0.12em; opacity: 0.78; font-weight: 700; }
      .sidebar-title { margin-top: 0.55rem; font-size: 1.25rem; font-weight: 800; line-height: 1.25; }
      .sidebar-meta { margin-top: 0.85rem; font-size: 0.9rem; color: rgba(255,255,255,0.88); }
      .toc { padding: 1rem 1rem 1.1rem; border-top: 0; background: linear-gradient(180deg, #f8fbff 0%, #f1f5f9 100%); }
      .toc h2 { margin: 0 0 0.9rem; font-size: 1.02rem; }
      .toc-list { list-style: none; margin: 0; padding: 0; display: block; }
      .toc-list li { margin: 0; }
      .toc-list a { display: block; color: #0f4c81; text-decoration: none; font-weight: 600; padding: 0.34rem 0.2rem; border-radius: 8px; }
      .toc-list a:hover { text-decoration: underline; }
      .content-panel { background: #ffffff; border-radius: 28px; padding: 2.5rem 2.75rem; box-shadow: 0 24px 60px rgba(15, 23, 42, 0.10); }
      h1, h2 { color: #0f172a; }
      h3 { color: #1e293b; margin-bottom: 0.4rem; }
      .grid { display: grid; grid-template-columns: repeat(auto-fit, minmax(220px, 1fr)); gap: 1rem; margin-bottom: 1.5rem; }
      .figure-grid { display: grid; grid-template-columns: repeat(auto-fit, minmax(320px, 1fr)); gap: 1.25rem; margin-top: 1rem; }
      .card { border: 1px solid #cbd5e1; border-radius: 14px; padding: 1rem; background: #f8fafc; }
      .figure-card { border: 1px solid #cbd5e1; border-radius: 16px; padding: 1rem; background: white; box-shadow: 0 12px 30px rgba(15, 23, 42, 0.06); }
      .hero { padding-bottom: 1.5rem; border-bottom: 1px solid #e2e8f0; margin-bottom: 1.75rem; }
      .hero-meta { color: #475569; margin-top: 0.5rem; }
      .metric { font-size: 1.7rem; font-weight: 700; }
      .section { margin: 2rem 0; }
      .note { color: #475569; }
      .eyebrow { font-size: 0.78rem; text-transform: uppercase; letter-spacing: 0.08em; color: #64748b; margin-bottom: 0.25rem; }
      table { width: 100%; border-collapse: collapse; margin-top: 1rem; font-size: 0.95rem; }
      th, td { border: 1px solid #e2e8f0; padding: 0.55rem; text-align: left; }
      th { background: #e2e8f0; }
      .card, .figure-card, p, li, td, th, code { overflow-wrap: anywhere; word-break: break-word; }
      img { width: 100%; border: 1px solid #e2e8f0; border-radius: 12px; margin: 0.75rem 0 0.5rem; background: #fff; }
      code { background: #e2e8f0; padding: 0.15rem 0.35rem; border-radius: 6px; }
      @media (max-width: 1100px) {
        .layout { grid-template-columns: 1fr; }
        .sidebar { position: static; }
      }
    </style>
  </head>
  <body>
    <div class="page-shell">
      <div class="layout">
        <aside class="sidebar">
          <div class="sidebar-hero">
            {% if atom_logo_src or entity_logo_src %}
            <div class="sidebar-logo-row">
              {% if atom_logo_src %}
              <img class="sidebar-logo" src="{{ atom_logo_src }}" alt="Logo ATOM" />
              {% endif %}
              {% if entity_logo_src %}
              <img class="sidebar-logo" src="{{ entity_logo_src }}" alt="Logo entidad" />
              {% endif %}
            </div>
            {% endif %}
            <p class="sidebar-eyebrow">Análisis</p>
            <div class="sidebar-title">Análisis de RNA-seq</div>
            {% if report_entity_name %}
            <p class="sidebar-meta">{{ report_entity_name }}</p>
            {% endif %}
            <p class="sidebar-meta"><strong>{{ report_author }}</strong></p>
            <p class="sidebar-meta">{{ report_design_id }} | {{ generated_at }}</p>
          </div>
          <nav class="toc" aria-label="Report index">
            <h2>Índice</h2>
            <ol class="toc-list">
              {% for section in report_sections %}
              <li><a href="#{{ section.id }}">{{ section.title }}</a></li>
              {% endfor %}
            </ol>
          </nav>
        </aside>
        <main class="content-panel">
    <div class="hero">
    <h1>Análisis de RNA-seq</h1>
    <p class="hero-meta"><strong>{{ report_design_id }}</strong> | {{ generated_at }}</p>
    <p>Contraste evaluado: <strong>{{ case }}</strong> frente a <strong>{{ control }}</strong> usando columna experimental <code>{{ condition_column }}</code>.</p>
    </div>

    <div class="grid">
      <div class="card"><div>Genes de entrada</div><div class="metric">{{ summary.n_input_genes }}</div></div>
      <div class="card"><div>Muestras de entrada</div><div class="metric">{{ summary.n_input_samples }}</div></div>
      <div class="card"><div>Genes evaluados</div><div class="metric">{{ summary.n_tested_genes }}</div></div>
      <div class="card"><div>Genes significativos</div><div class="metric">{{ summary.n_significant_genes }}</div></div>
      <div class="card"><div>Sobreexpresados</div><div class="metric">{{ summary.n_upregulated }}</div></div>
      <div class="card"><div>Infraexpresados</div><div class="metric">{{ summary.n_downregulated }}</div></div>
    </div>

    <div class="section" id="section-introduccion">
    <h2>0.1 Introducción</h2>
    <p>{{ intro_text }}</p>
    <p>
      Este flujo en Python replica, en la medida de lo posible, pasos principales del pipeline en R:
      control de calidad, expresión diferencial, visualización multigráfico y enriquecimiento funcional.
    </p>
    </div>

    <div class="section">
    <h2>Resumen ejecutivo</h2>
    <p>
      Este pipeline en Python evaluó <strong>{{ case }}</strong> frente a <strong>{{ control }}</strong> sobre
      <strong>{{ summary.n_input_samples }}</strong> muestras seleccionadas. La expresión diferencial identificó
      <strong>{{ summary.n_significant_genes }}</strong> genes que superaron umbrales configurados
      (adjusted p-value &lt; {{ summary.alpha }} y log2 fold change absoluto &ge; {{ summary.log2fc_threshold }}).
      Los cambios observados oscilaron entre <strong>{{ summary.min_log2fc }}</strong> y <strong>{{ summary.max_log2fc }}</strong>.
    </p>
    </div>

    <div class="section" id="section-sample-distribution">
    <h2>0.2 Información de las muestras</h2>
    <div class="grid">
      {% for item in sample_distribution %}
      <div class="card">
        <div>{{ item.group }}</div>
        <div class="metric">{{ item.count }}</div>
      </div>
      {% endfor %}
    </div>
    <p class="note">
      Resumen de control de calidad por muestra en secciones de PCA y heatmap mostradas más abajo.
    </p>
    </div>

    <div class="section" id="section-quality-control">
    <h2>0.3 Control de calidad</h2>
    <p>{{ qc_summary }}</p>
    </div>

    <div class="section" id="section-de-overview">
    <h2>0.4 Expresión Génica Diferencial</h2>
    <div class="grid">
      <div class="card">
        <h3>Genes más sobreexpresados</h3>
        {% if top_up_genes %}
        <p>{{ top_up_genes | join(", ") }}</p>
        {% else %}
        <p>Ningún gen sobreexpresado superó umbrales configurados.</p>
        {% endif %}
      </div>
      <div class="card">
        <h3>Genes más infraexpresados</h3>
        {% if top_down_genes %}
        <p>{{ top_down_genes | join(", ") }}</p>
        {% else %}
        <p>Ningún gen infraexpresado superó umbrales configurados.</p>
        {% endif %}
      </div>
    </div>
    <div class="grid">
      <div class="card">
        <h3>Destacados sobreexpresados</h3>
        <ul>
          {% for item in top_up_details %}
          <li>{{ item }}</li>
          {% endfor %}
        </ul>
      </div>
      <div class="card">
        <h3>Destacados infraexpresados</h3>
        <ul>
          {% for item in top_down_details %}
          <li>{{ item }}</li>
          {% endfor %}
        </ul>
      </div>
    </div>
    <p class="note">
      Tabla de genes destacados conserva valores estadísticos crudos para comparación directa frente a pipeline en R.
    </p>
    </div>

    <div class="section" id="section-top-hits">
    <h2>0.5 Genes destacados</h2>
    <table>
      <thead>
        <tr>
          {% for column in ["gene_symbol", "gene_id", "log2FoldChange", "pvalue", "padj"] %}
          <th>{{ column }}</th>
          {% endfor %}
        </tr>
      </thead>
      <tbody>
        {% for row in top_hits %}
        <tr>
          <td>{{ row.get("gene_symbol", "") }}</td>
          <td>{{ row.get("gene_id", "") }}</td>
          <td>{{ row.get("log2FoldChange", "") }}</td>
          <td>{{ row.get("pvalue", "") }}</td>
          <td>{{ row.get("padj", "") }}</td>
        </tr>
        {% endfor %}
      </tbody>
    </table>
    </div>

    <div class="section" id="section-figure-overview">
    <h2>0.6 Resumen de figuras</h2>
    <div class="grid">
      <div class="card"><div>Figuras QC</div><div class="metric">{{ qc_images | length }}</div></div>
      <div class="card"><div>Figuras DEG</div><div class="metric">{{ de_images | length }}</div></div>
      <div class="card"><div>Figuras enriquecimiento</div><div class="metric">{{ enrichment_images | length }}</div></div>
      <div class="card"><div>Total figuras</div><div class="metric">{{ image_entries | length }}</div></div>
    </div>
    </div>

    <div class="section" id="section-qc-figures">
    <h2>0.7 Figuras de control de calidad</h2>
    <div class="figure-grid">
      {% for image in qc_images %}
      <div class="figure-card">
        <div class="eyebrow">{{ image.group }}</div>
        <h3>{{ image.label }}</h3>
        <p class="note">{{ image.caption }}</p>
        <img src="{{ image.src }}" alt="{{ image.label }}" />
      </div>
      {% endfor %}
    </div>
    </div>

    <div class="section" id="section-de-figures">
    <h2>0.8 Figuras de expresión diferencial</h2>
    <div class="figure-grid">
      {% for image in de_images %}
      <div class="figure-card">
        <div class="eyebrow">{{ image.group }}</div>
        <h3>{{ image.label }}</h3>
        <p class="note">{{ image.caption }}</p>
        <img src="{{ image.src }}" alt="{{ image.label }}" />
      </div>
      {% endfor %}
    </div>
    {% if other_images %}
    <div class="figure-grid">
      {% for image in other_images %}
      <div class="figure-card">
        <div class="eyebrow">{{ image.group }}</div>
        <h3>{{ image.label }}</h3>
        <p class="note">{{ image.caption }}</p>
        <img src="{{ image.src }}" alt="{{ image.label }}" />
      </div>
      {% endfor %}
    </div>
    {% endif %}
    </div>

    <div class="section" id="section-enrichment-figures">
    <h2>0.9 Figuras de enriquecimiento funcional</h2>
    <div class="figure-grid">
      {% for image in enrichment_images %}
      <div class="figure-card">
        <div class="eyebrow">{{ image.group }}</div>
        <h3>{{ image.label }}</h3>
        <p class="note">{{ image.caption }}</p>
        <img src="{{ image.src }}" alt="{{ image.label }}" />
      </div>
      {% endfor %}
    </div>
    </div>

    <div class="section" id="section-functional-enrichment">
    <h2>0.10 Enriquecimiento funcional</h2>
    <div class="grid">
      <div class="card">
        <h3>Enriquecimiento prerankeado</h3>
        <p>{{ enrichment_summary.prerank }}</p>
      </div>
      <div class="card">
        <h3>Programa sobreexpresado</h3>
        <p>{{ enrichment_summary.up }}</p>
      </div>
      <div class="card">
        <h3>Programa infraexpresado</h3>
        <p>{{ enrichment_summary.down }}</p>
      </div>
    </div>
    </div>

    <div class="section" id="section-interpretation">
    <h2>0.11 Interpretación biológica integrada</h2>
    <p>{{ integrated_interpretation }}</p>
    </div>

    <div class="section" id="section-validation">
    <h2>0.12 Validación o siguientes pasos</h2>
    <ul>
      {% for item in validation_suggestions %}
      <li>{{ item }}</li>
      {% endfor %}
    </ul>
    </div>

    <div class="section" id="section-limitations">
    <h2>0.13 Limitaciones</h2>
    <ul>
      {% for item in limitations %}
      <li>{{ item }}</li>
      {% endfor %}
    </ul>
    </div>

    <div class="section" id="section-enrichment-summary">
    <h2>0.14 Resumen de enriquecimiento</h2>
    {% if enrichment_overview %}
    <ul>
      {% for name, row_count in enrichment_overview.items() %}
      <li><strong>{{ name }}</strong>: {{ row_count }} filas</li>
      {% endfor %}
    </ul>
    {% for name, rows in enrichment_highlights.items() %}
    <h3>{{ name }}</h3>
    <table>
      <thead>
        <tr><th>término</th><th>score</th></tr>
      </thead>
      <tbody>
        {% for row in rows %}
        <tr><td>{{ row.term }}</td><td>{{ row.score }}</td></tr>
        {% endfor %}
      </tbody>
    </table>
    {% endfor %}
    {% else %}
    <p>No se generó salida de enriquecimiento. Causas comunes: anotación de símbolos incompleta, conjuntos pequeños o señal biológica demasiado estrecha para umbrales actuales.</p>
    {% endif %}
    </div>

    <div class="section" id="section-references">
    <h2>0.15 Referencias</h2>
    <ul>
      {% for item in curated_references %}
      <li>{{ item }}</li>
      {% endfor %}
    </ul>
    </div>
        </main>
      </div>
    </div>
  </body>
</html>
        """
    )

    html_path = output_dir / "report.html"
    html_path.write_text(
        template.render(
            case=case,
            control=control,
            condition_column=condition_column,
            summary=summary,
            top_hits=top_hits_records,
            sample_distribution=sample_distribution,
            image_entries=image_entries,
            qc_images=qc_images,
            de_images=de_images,
            enrichment_images=enrichment_images,
            other_images=other_images,
            enrichment_overview=enrichment_overview,
            enrichment_highlights=enrichment_highlights,
            enrichment_summary=enrichment_summary,
            qc_summary=qc_summary,
            integrated_interpretation=integrated_interpretation,
            validation_suggestions=validation_suggestions,
            limitations=limitations,
            curated_references=curated_references,
            top_up_genes=top_up_genes,
            top_down_genes=top_down_genes,
            top_up_details=top_up_details,
            top_down_details=top_down_details,
            report_sections=report_sections,
            report_design_id=report_design_id,
            generated_at=generated_at,
            report_author=report_author,
            report_entity_name=report_entity_name,
            atom_logo_src=atom_logo_src,
            entity_logo_src=entity_logo_src,
            intro_text=intro_text,
        ),
        encoding="utf-8",
    )
    LOGGER.info("Generated HTML report at %s", html_path)
    return html_path


def _generate_docx_report(
    output_dir: Path,
    de_outputs: DifferentialExpressionOutputs,
    generated_assets: dict[str, Path | None],
    enrichment_results: dict[str, Any],
    condition_column: str,
    case: str,
    control: str,
    branding: ReportBranding,
    report_title: str,
    report_subtitle: str,
    report_organism: str,
    report_molecule: str,
    report_cell_context: str,
    report_condition_description: str,
    report_mode_label: str,
) -> Path:
    docx_module = _import_dependency("docx", "python-docx")

    document = docx_module.Document()
    _configure_document_styles(document, docx_module)
    _configure_docx_header(document, output_dir, branding, docx_module)
    logo_paragraph = document.add_paragraph()
    logo_paragraph.alignment = docx_module.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    added_logo = False
    atom_logo_path = _resolve_atom_logo_path()
    safe_atom_logo_path = _prepare_docx_image_path(atom_logo_path, output_dir) if atom_logo_path else None
    if safe_atom_logo_path and safe_atom_logo_path.exists():
        logo_run = logo_paragraph.add_run()
        logo_run.add_picture(str(safe_atom_logo_path), width=docx_module.shared.Inches(0.9))
        added_logo = True
    if branding.entity_logo_path:
        logo_path = Path(branding.entity_logo_path)
        safe_logo_path = _prepare_docx_image_path(logo_path, output_dir)
        if safe_logo_path and safe_logo_path.exists():
            if added_logo:
                logo_paragraph.add_run("   ")
            logo_run = logo_paragraph.add_run()
            logo_run.add_picture(str(safe_logo_path), width=docx_module.shared.Inches(0.9))
            added_logo = True
    if not added_logo:
        document.add_paragraph()
    document.add_paragraph(report_title, style="Title")
    subtitle = document.add_paragraph(style="Subtitle")
    subtitle.paragraph_format.space_after = 0
    report_design_id = output_dir.name.split("__")[0] if "__" in output_dir.name else output_dir.name
    generated_at = datetime.now().strftime("%Y-%m-%d %H:%M")
    subtitle_lines = []
    if report_subtitle:
        subtitle_lines.append(report_subtitle)
    if branding.entity_name:
        subtitle_lines.append(branding.entity_name)
    subtitle_lines.append(branding.author)
    subtitle_lines.append(generated_at)
    subtitle.add_run("\n".join(subtitle_lines))
    _add_docx_divider(document, docx_module)

    _add_docx_section_heading(document, "Study Snapshot", level=1)
    snapshot_items = [
        f"Design ID: {report_design_id}",
        f"Entity: {branding.entity_name}" if branding.entity_name else "",
        f"Project owner: {branding.author}" if branding.author else "",
        f"Contrast: {control} vs {case}",
        f"Organism: {report_organism}" if report_organism else "",
        f"Molecule: {report_molecule}" if report_molecule else "",
        f"Cell context: {report_cell_context}" if report_cell_context else "",
        f"Condition: {report_condition_description}" if report_condition_description else "",
        f"Report mode: {report_mode_label}" if report_mode_label else "",
    ]
    for item in snapshot_items:
        if item:
            _add_docx_body_paragraph(document, item, style="List Bullet")
    _add_docx_divider(document, docx_module)

    summary = de_outputs.summary
    sample_distribution = _build_sample_distribution(de_outputs.metadata, condition_column)
    enrichment_highlights = _build_enrichment_highlights(enrichment_results)
    enrichment_summary = _build_functional_enrichment_summary(enrichment_highlights)
    top_up_genes = [str(gene) for gene in summary.get("top_upregulated_genes", [])[:10]]
    top_down_genes = [str(gene) for gene in summary.get("top_downregulated_genes", [])[:10]]
    top_up_details = _build_ranked_gene_details(de_outputs.upregulated.nsmallest(6, "padj"), limit=6)
    top_down_details = _build_ranked_gene_details(de_outputs.downregulated.nsmallest(6, "padj"), limit=6)
    qc_summary = _build_qc_summary(sample_distribution)
    integrated_interpretation = _build_integrated_interpretation(
        top_up_genes,
        top_down_genes,
        enrichment_summary,
    )
    validation_suggestions = _build_validation_suggestions(top_up_genes, top_down_genes)
    limitations = _build_limitations(summary, enrichment_highlights)
    curated_references = _build_python_curated_references()
    image_entries = _build_image_entries(generated_assets)
    qc_images = [entry for entry in image_entries if entry["group"] == "Control de calidad"]
    de_images = [entry for entry in image_entries if entry["group"] == "Expresión diferencial"]
    enrichment_images = [entry for entry in image_entries if entry["group"] == "Enriquecimiento funcional"]

    _add_docx_metric_table(document, summary, docx_module)
    _add_docx_divider(document, docx_module)

    _add_docx_section_heading(document, "Executive Summary", level=1)
    _add_docx_body_paragraph(
        document,
        " ".join(
            [
                f"This study evaluated {control} versus {case} across {summary.get('n_input_samples')} samples.",
                f"Differential expression identified {summary.get('n_significant_genes')} significant genes",
                f"with {summary.get('n_upregulated')} upregulated and {summary.get('n_downregulated')} downregulated.",
                f"Observed log2 fold changes ranged from {_safe_float(summary.get('min_log2fc'))} to {_safe_float(summary.get('max_log2fc'))}.",
            ]
        ),
    )

    _add_docx_section_heading(document, "Quality Control and Sample Separation", level=1)
    for item in sample_distribution:
        _add_docx_body_paragraph(document, f"{item['group']}: {item['count']}", style="List Bullet")
    _add_docx_body_paragraph(document, qc_summary)

    _add_docx_section_heading(document, "Differential Expression Overview", level=1)
    if top_up_genes:
        _add_docx_body_paragraph(document, "Top upregulated genes: " + ", ".join(top_up_genes[:8]))
    if top_down_genes:
        _add_docx_body_paragraph(document, "Top downregulated genes: " + ", ".join(top_down_genes[:8]))
    if top_up_details:
        _add_docx_section_heading(document, "Upregulated highlights", level=2)
        for item in top_up_details:
            _add_docx_body_paragraph(document, item, style="List Bullet")
    if top_down_details:
        _add_docx_section_heading(document, "Downregulated highlights", level=2)
        for item in top_down_details:
            _add_docx_body_paragraph(document, item, style="List Bullet")

    _add_docx_section_heading(document, "Key significant genes", level=1)
    for item in _build_ranked_gene_details(de_outputs.significant.nsmallest(10, "padj"), limit=10):
        _add_docx_body_paragraph(document, item, style="List Bullet")

    if qc_images:
        _add_docx_section_heading(document, "Quality Control Figures", level=1)
        for entry in qc_images:
            _add_docx_figure(document, _resolve_figure_key_from_label(entry["label"]), output_dir / entry["filename"], docx_module)
        _add_docx_divider(document, docx_module)

    if de_images:
        _add_docx_section_heading(document, "Differential Expression Figures", level=1)
        for entry in de_images:
            _add_docx_figure(document, _resolve_figure_key_from_label(entry["label"]), output_dir / entry["filename"], docx_module)
        _add_docx_divider(document, docx_module)

    if enrichment_images:
        _add_docx_section_heading(document, "Functional Enrichment Figures", level=1)
        for entry in enrichment_images:
            _add_docx_figure(document, _resolve_figure_key_from_label(entry["label"]), output_dir / entry["filename"], docx_module)
        _add_docx_divider(document, docx_module)

    _add_docx_section_heading(document, "Functional Enrichment", level=1)
    _add_docx_body_paragraph(document, "Preranked enrichment: " + enrichment_summary["prerank"])
    _add_docx_body_paragraph(document, "Upregulated program: " + enrichment_summary["up"])
    _add_docx_body_paragraph(document, "Downregulated program: " + enrichment_summary["down"])

    _add_docx_section_heading(document, "Integrated Biological Interpretation", level=1)
    _add_docx_body_paragraph(document, integrated_interpretation)

    _add_docx_section_heading(document, "Suggested Validation or Follow-up", level=1)
    for item in validation_suggestions:
        _add_docx_body_paragraph(document, item, style="List Bullet")

    _add_docx_section_heading(document, "Limitations", level=1)
    for item in limitations:
        _add_docx_body_paragraph(document, item, style="List Bullet")

    _add_docx_section_heading(document, "Enrichment Highlights", level=1)
    if enrichment_highlights:
        for name, rows in enrichment_highlights.items():
            _add_docx_section_heading(document, f"{name}", level=2)
            for row in rows:
                score_suffix = f" ({row['score']})" if row["score"] else ""
                _add_docx_body_paragraph(document, f"{row['term']}{score_suffix}", style="List Bullet")
    else:
        _add_docx_body_paragraph(
            document,
            "No enrichment output was produced. Likely causes include incomplete symbol annotation, small gene sets, or low overlap with configured libraries."
        )

    _add_docx_section_heading(document, "References", level=1)
    for item in curated_references:
        _add_docx_body_paragraph(document, item, style="List Bullet")

    output_path = output_dir / "report.docx"
    document.save(output_path)
    LOGGER.info("Generated DOCX report at %s", output_path)
    return output_path


def create_zip_package(output_dir: Path, zip_name: str = "rna_seq_python_results.zip") -> Path:
    zip_path = output_dir / zip_name
    with ZipFile(zip_path, mode="w", compression=ZIP_DEFLATED) as archive:
        for path in sorted(output_dir.rglob("*")):
            if not path.is_file() or path == zip_path:
                continue
            archive.write(path, arcname=path.relative_to(output_dir))
    LOGGER.info("Created ZIP package at %s", zip_path)
    return zip_path


def _build_arg_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="Python RNA-seq differential expression and enrichment pipeline.")
    parser.add_argument("--counts", required=True, help="Counts matrix file path (.csv, .tsv, .txt, .xlsx).")
    parser.add_argument("--metadata", required=True, help="Metadata file path (.csv, .tsv, .txt, .xlsx).")
    parser.add_argument("--condition-column", required=True, help="Metadata column defining contrast groups.")
    parser.add_argument("--case", required=True, help="Case group label.")
    parser.add_argument("--control", required=True, help="Control group label.")
    parser.add_argument("--output-dir", required=True, help="Directory where results will be written.")
    parser.add_argument("--report-author", default="Juan Vladimir de la Rosa Medina", help="Author or owner label for report header.")
    parser.add_argument("--report-entity-name", default="", help="Optional entity name for report header.")
    parser.add_argument("--report-entity-logo-path", default="", help="Optional local entity logo path for report header.")
    parser.add_argument("--report-title", default="RNA-seq Python Report", help="Report title for DOCX/HTML.")
    parser.add_argument("--report-subtitle", default="", help="Report subtitle line.")
    parser.add_argument("--report-organism", default="", help="Organism label for report snapshot.")
    parser.add_argument("--report-molecule", default="", help="Molecule label for report snapshot.")
    parser.add_argument("--report-cell-context", default="", help="Cell context label for report snapshot.")
    parser.add_argument("--report-condition-description", default="", help="Condition description for report snapshot.")
    parser.add_argument("--report-mode-label", default="Python workflow", help="Report mode label for report snapshot.")
    parser.add_argument("--alpha", type=float, default=0.05, help="Adjusted p-value threshold. Default: 0.05.")
    parser.add_argument(
        "--log2fc-threshold",
        type=float,
        default=1.0,
        help="Absolute log2 fold change threshold for significance. Default: 1.0.",
    )
    parser.add_argument(
        "--top-heatmap-genes",
        type=int,
        default=50,
        help="Maximum number of significant genes to draw in heatmap. Default: 50.",
    )
    parser.add_argument(
        "--docx",
        action="store_true",
        help="Also generate DOCX report using python-docx.",
    )
    parser.add_argument(
        "--verbose",
        action="store_true",
        help="Enable debug logging.",
    )
    return parser


def main() -> int:
    parser = _build_arg_parser()
    args = parser.parse_args()
    _configure_logging(verbose=args.verbose)

    output_dir = Path(args.output_dir).expanduser().resolve()
    output_dir.mkdir(parents=True, exist_ok=True)
    branding = _resolve_report_branding(args.report_author, args.report_entity_name, args.report_entity_logo_path)

    try:
        loaded = load_inputs(args.counts, args.metadata)
        counts, metadata = validate_inputs(
            loaded.counts,
            loaded.metadata,
            args.condition_column,
            args.case,
            args.control,
            loaded.sample_column,
        )
        de_outputs = run_differential_expression(
            counts,
            metadata,
            args.condition_column,
            args.case,
            args.control,
            alpha=args.alpha,
            log2fc_threshold=args.log2fc_threshold,
        )

        generated_assets: dict[str, Path | None] = {}
        generated_assets["pca"] = generate_pca(
            de_outputs.normalized_counts,
            de_outputs.metadata,
            args.condition_column,
            output_dir,
        )
        generated_assets["sample_distance_heatmap"] = generate_sample_distance_heatmap(
            de_outputs.normalized_counts,
            de_outputs.metadata,
            args.condition_column,
            output_dir,
        )
        generated_assets["library_size_barplot"] = generate_library_size_barplot(
            de_outputs.counts,
            de_outputs.metadata,
            args.condition_column,
            output_dir,
        )
        generated_assets["normalized_expression_boxplot"] = generate_normalized_expression_boxplot(
            de_outputs.normalized_counts,
            de_outputs.metadata,
            args.condition_column,
            output_dir,
        )
        generated_assets["volcano_plot"] = generate_volcano_plot(
            de_outputs.results,
            output_dir,
            alpha=args.alpha,
            log2fc_threshold=args.log2fc_threshold,
        )
        generated_assets["ma_plot"] = generate_ma_plot(
            de_outputs.results,
            output_dir,
            alpha=args.alpha,
        )
        generated_assets["heatmap"] = generate_heatmap(
            de_outputs.normalized_counts,
            de_outputs.significant,
            de_outputs.metadata,
            args.condition_column,
            output_dir,
            top_n=args.top_heatmap_genes,
        )
        generated_assets["top_significant_genes_barplot"] = generate_top_genes_barplot(
            de_outputs.significant,
            output_dir,
        )
        generated_assets["pvalue_histogram"] = generate_pvalue_histogram(
            de_outputs.results,
            output_dir,
        )

        enrichment_results = run_enrichment(de_outputs, output_dir)
        generated_assets["gsea_prerank_dotplot"] = generate_enrichment_dotplot(
            enrichment_results.get("gsea_prerank"),
            output_dir,
            "gsea_prerank_dotplot.png",
            "Preranked GSEA overview",
        )
        generated_assets["enrichr_up_dotplot"] = generate_enrichment_dotplot(
            enrichment_results.get("enrichr_up"),
            output_dir,
            "enrichr_up_dotplot.png",
            "Upregulated enrichment overview",
        )
        generated_assets["enrichr_down_dotplot"] = generate_enrichment_dotplot(
            enrichment_results.get("enrichr_down"),
            output_dir,
            "enrichr_down_dotplot.png",
            "Downregulated enrichment overview",
        )
        exported_paths = export_results(de_outputs, enrichment_results, output_dir)
        html_report = generate_html_report(
            output_dir,
            de_outputs,
            generated_assets,
            enrichment_results,
            args.condition_column,
            args.case,
            args.control,
            branding,
        )
        exported_paths["html_report"] = html_report

        if args.docx:
            exported_paths["docx_report"] = _generate_docx_report(
                output_dir,
                de_outputs,
                generated_assets,
                enrichment_results,
                args.condition_column,
                args.case,
                args.control,
                branding,
                args.report_title,
                args.report_subtitle,
                args.report_organism,
                args.report_molecule,
                args.report_cell_context,
                args.report_condition_description,
                args.report_mode_label,
            )

        de_outputs.summary["generated_files"] = {
            key: str(path.relative_to(output_dir))
            for key, path in exported_paths.items()
            if path.exists()
        }
        _save_summary_json(output_dir, de_outputs.summary)
        exported_paths["zip_package"] = create_zip_package(output_dir)
        LOGGER.info("Pipeline finished successfully.")
        return 0
    except Exception as exc:
        LOGGER.exception("Pipeline failed: %s", exc)
        return 1


if __name__ == "__main__":
    sys.exit(main())
