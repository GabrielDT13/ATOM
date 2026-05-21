from __future__ import annotations

import csv
import json
import shutil
import subprocess
import sys
from datetime import UTC, datetime
from pathlib import Path
from typing import Iterator

from backend.app.core.config import get_settings
from backend.app.services.dashboard_activity import log_analysis_dashboard_event
from backend.app.services.entities import get_entity_logo_path
from backend.app.services.project_inventory import read_project_settings
from backend.app.services.project_repository import _get_project_record
from backend.app.services.projects import get_project_dir
from backend.app.services.users import _get_profile_by_username
from openpyxl import load_workbook


DEFAULT_REPORT_AUTHOR = "Juan Vladimir de la Rosa Medina"


def _stream_event(payload: dict[str, object]) -> str:
    return f"data:{json.dumps(payload, ensure_ascii=False)}\n\n"


def _timestamp() -> str:
    return datetime.now(UTC).isoformat()


def parse_analysis_variant_from_trigger_source(trigger_source: str | None) -> str | None:
    normalized = str(trigger_source or "").strip().lower()
    parts = [part.strip() for part in normalized.split(":") if part.strip()]
    if len(parts) < 2:
        return None
    variant = parts[1]
    if variant in {"basic", "enhanced", "python"}:
        return variant
    return None


def resolve_analysis_variant(
    project_dir: Path,
    analysis_type: str,
    *,
    preferred_variant: str | None = None,
) -> str:
    normalized_analysis_type = str(analysis_type or "").strip().lower()
    if normalized_analysis_type != "rna-seq":
        return "basic"

    project_settings = read_project_settings(project_dir)
    enabled_variants = {
        str(item or "").strip().lower()
        for item in project_settings.get("enabled_analysis_variants", [])
        if str(item or "").strip()
    }
    if not enabled_variants:
        enabled_variants = {"basic"}

    normalized_preferred_variant = str(preferred_variant or "").strip().lower()
    if normalized_preferred_variant in enabled_variants:
        return normalized_preferred_variant

    primary_variant = str(project_settings.get("primary_analysis_variant") or "").strip().lower()
    if primary_variant in enabled_variants:
        return primary_variant

    if "basic" in enabled_variants:
        return "basic"
    return sorted(enabled_variants)[0]


def resolve_analysis_script_key(
    project_dir: Path,
    analysis_type: str,
    *,
    preferred_variant: str | None = None,
) -> str:
    normalized_analysis_type = str(analysis_type or "").strip().lower()
    if normalized_analysis_type != "rna-seq":
        return normalized_analysis_type

    variant = resolve_analysis_variant(
        project_dir,
        analysis_type,
        preferred_variant=preferred_variant,
    )
    if variant == "python":
        return "rna-seq-python"
    if variant == "enhanced":
        return "rna-seq-pro"
    return "rna-seq"


def resolve_design_output_dir(project_dir: Path, design_id: str, script_key: str) -> Path:
    normalized_design_id = str(design_id or "").strip()
    normalized_script_key = str(script_key or "").strip()
    if not normalized_design_id:
        raise ValueError("design_id es obligatorio")
    if not normalized_script_key:
        return project_dir / normalized_design_id
    return project_dir / f"{normalized_design_id}__{normalized_script_key}"


def clean_resultados(output_dir: Path, design_id: str) -> None:
    keep_ext = {".docx", ".html", ".json", ".png", ".rmd", ".tsv", ".txt", ".xlsx", ".zip"}
    keep_files = {
        output_dir / f"{design_id}.docx",
        output_dir / f"{design_id}.html",
        output_dir / f"{design_id}.Rmd",
        output_dir / f"{design_id}.xlsx",
        output_dir / f"{design_id}.zip",
        output_dir / "differential_expression_results.tsv",
        output_dir / "pca_samples.png",
        output_dir / "python_metadata.tsv",
        output_dir / "significant_genes.tsv",
        output_dir / "significant_genes_heatmap.png",
        output_dir / "summary_metrics.json",
        output_dir / "volcano_plot.png",
    }

    for item in output_dir.iterdir():
        if item in keep_files or item.suffix.lower() in keep_ext:
            continue
        if item.is_dir():
            shutil.rmtree(item)
        else:
            item.unlink()


def reset_output_dir(output_dir: Path) -> None:
    if not output_dir.exists():
        return

    for item in output_dir.iterdir():
        if item.is_dir():
            shutil.rmtree(item)
        else:
            item.unlink()


def _resolve_atom_logo_path(settings) -> Path | None:
    candidate = settings.project_root / "frontend" / "public" / "images" / "logo.png"
    if candidate.exists() and candidate.is_file():
        return candidate
    return None


def _prepare_docx_image_path(path: Path | None, output_dir: Path) -> Path | None:
    if path is None or not path.exists() or not path.is_file():
        return None

    if path.suffix.lower() in {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tif", ".tiff"}:
        return path

    try:
        from PIL import Image

        converted_path = output_dir / f"{path.stem}_header_docx.png"
        with Image.open(path) as image:
            if image.mode not in {"RGB", "RGBA"}:
                image = image.convert("RGBA")
            image.save(converted_path, format="PNG")
        return converted_path
    except Exception:
        return None


def _apply_docx_branding(docx_path: Path, output_dir: Path, report_branding: dict[str, str], settings) -> None:
    if not docx_path.exists() or not docx_path.is_file():
        return

    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches
    except Exception:
        return

    atom_logo_path = _prepare_docx_image_path(_resolve_atom_logo_path(settings), output_dir)
    entity_logo_value = str(report_branding.get("report_entity_logo_path") or "").strip()
    entity_logo_path = _prepare_docx_image_path(Path(entity_logo_value), output_dir) if entity_logo_value else None
    if not atom_logo_path and not entity_logo_path:
        return

    document = Document(str(docx_path))
    for section in document.sections:
        header = section.header
        header.is_linked_to_previous = False
        paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        if paragraph.runs:
            for run in list(paragraph.runs):
                run.clear()
        paragraph.text = ""
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if atom_logo_path and atom_logo_path.exists():
            paragraph.add_run().add_picture(str(atom_logo_path), width=Inches(0.55))
        if entity_logo_path and entity_logo_path.exists():
            if atom_logo_path and atom_logo_path.exists():
                paragraph.add_run("   ")
            paragraph.add_run().add_picture(str(entity_logo_path), width=Inches(0.55))
    document.save(str(docx_path))


def resolve_template_file(project_dir: Path) -> Path:
    preferred = project_dir / "template.xlsx"
    if preferred.exists():
        return preferred

    for candidate in sorted(project_dir.iterdir()):
        if candidate.is_file() and candidate.suffix.lower() in {".xlsx", ".xls"}:
            return candidate
    raise FileNotFoundError("Archivo template.xlsx no encontrado")


def _worksheet_records(workbook, sheet_name: str) -> list[dict[str, object]]:
    if sheet_name not in workbook.sheetnames:
        raise ValueError(f"Hoja '{sheet_name}' no encontrada en template.xlsx")

    sheet = workbook[sheet_name]
    try:
        headers = [str(cell).strip() if cell is not None else "" for cell in next(sheet.iter_rows(values_only=True))]
    except StopIteration as exc:
        raise ValueError(f"La hoja '{sheet_name}' está vacía") from exc

    return [dict(zip(headers, row)) for row in sheet.iter_rows(min_row=2, values_only=True)]


def _string_value(value: object) -> str:
    return str(value or "").strip()


def _filter_metadata_records(
    metadata_rows: list[dict[str, object]],
    design_row: dict[str, object],
) -> list[dict[str, object]]:
    filters = {
        "treatment": {_string_value(design_row.get("treatment1")), _string_value(design_row.get("treatment2"))},
        "genotype": {_string_value(design_row.get("genotype1")), _string_value(design_row.get("genotype2"))},
        "gender": {_string_value(design_row.get("gender1")), _string_value(design_row.get("gender2"))},
        "cell": {_string_value(design_row.get("cell1")), _string_value(design_row.get("cell2"))},
    }

    filtered_rows = metadata_rows
    for key, values in filters.items():
        normalized_values = {item for item in values if item}
        if not normalized_values:
            continue
        filtered_rows = [row for row in filtered_rows if _string_value(row.get(key)) in normalized_values]

    if not filtered_rows:
        raise ValueError("Ninguna muestra de metadata coincide con filtros del diseño seleccionado")
    return filtered_rows


def _write_python_metadata_file(output_dir: Path, metadata_rows: list[dict[str, object]]) -> Path:
    if not metadata_rows:
        raise ValueError("No hay metadata para exportar")

    fieldnames = [str(key).strip() for key in metadata_rows[0].keys() if str(key).strip()]
    metadata_path = output_dir / "python_metadata.tsv"
    with metadata_path.open("w", encoding="utf-8", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=fieldnames, delimiter="\t", extrasaction="ignore")
        writer.writeheader()
        for row in metadata_rows:
            writer.writerow({field: row.get(field) for field in fieldnames})
    return metadata_path


def _resolve_python_counts_file(project_dir: Path, metadata_rows: list[dict[str, object]]) -> Path:
    count_files = sorted(
        {
            _string_value(row.get("counts_file"))
            for row in metadata_rows
            if _string_value(row.get("counts_file"))
        }
    )
    if not count_files:
        raise ValueError("Python RNA-seq requiere counts_file en metadata")
    if len(count_files) != 1:
        raise ValueError("Python RNA-seq requiere un único counts_file por diseño")

    count_path = project_dir / count_files[0]
    if not count_path.exists():
        raise FileNotFoundError(f"Archivo de counts no encontrado: {count_files[0]}")
    return count_path


def _resolve_python_contrast(
    metadata_rows: list[dict[str, object]],
    design_row: dict[str, object],
) -> tuple[str, str, str]:
    condition_column = _string_value(design_row.get("analysis_design")) or "condition"
    observed_values = sorted(
        {
            _string_value(row.get(condition_column))
            for row in metadata_rows
            if _string_value(row.get(condition_column))
        }
    )
    if len(observed_values) != 2:
        raise ValueError(
            f"Python RNA-seq requiere exactamente dos grupos en '{condition_column}', encontrados: {observed_values}"
        )
    control, case = observed_values[0], observed_values[1]
    return condition_column, case, control


def _sync_python_report_artifacts(output_dir: Path, design_id: str) -> None:
    artifact_pairs = {
        "report.html": f"{design_id}.html",
        "report.docx": f"{design_id}.docx",
        "rna_seq_python_results.xlsx": f"{design_id}.xlsx",
        "rna_seq_python_results.zip": f"{design_id}.zip",
    }
    for source_name, target_name in artifact_pairs.items():
        source_path = output_dir / source_name
        target_path = output_dir / target_name
        if not source_path.exists():
            continue
        shutil.copy2(source_path, target_path)
        if source_path != target_path and source_path.exists():
            source_path.unlink()


def _resolve_report_branding(project_owner_username: str, project_name: str) -> dict[str, str]:
    project_record = _get_project_record(project_owner_username, project_name) or {}
    owner_profile = _get_profile_by_username(project_owner_username) or {}

    owner_display_name = str(owner_profile.get("full_name") or "").strip()
    owner_username = str(owner_profile.get("username") or "").strip() or project_owner_username
    entity_name = str(project_record.get("entity_name") or "").strip()
    entity_id = str(project_record.get("entity_id") or "").strip()

    entity_logo_path = ""
    if entity_id:
        candidate_logo_path = get_entity_logo_path(entity_id)
        if candidate_logo_path.exists() and candidate_logo_path.is_file():
            entity_logo_path = str(candidate_logo_path)

    return {
        "report_author": owner_display_name or owner_username or DEFAULT_REPORT_AUTHOR,
        "report_entity_name": entity_name,
        "report_entity_logo_path": entity_logo_path,
    }


def _run_python_analysis(
    *,
    project_dir: Path,
    output_dir: Path,
    design_id: str,
    design_row: dict[str, object],
    metadata_rows: list[dict[str, object]],
    report_branding: dict[str, str],
) -> list[str]:
    settings = get_settings()
    counts_path = _resolve_python_counts_file(project_dir, metadata_rows)
    metadata_path = _write_python_metadata_file(output_dir, metadata_rows)
    condition_column, case, control = _resolve_python_contrast(metadata_rows, design_row)
    python_script_path = settings.project_root / "scripts" / "rna_seq_python_poc.py"
    if not python_script_path.exists():
        raise FileNotFoundError("scripts/rna_seq_python_poc.py no encontrado")

    command = [
        sys.executable,
        str(python_script_path),
        "--counts",
        str(counts_path),
        "--metadata",
        str(metadata_path),
        "--condition-column",
        condition_column,
        "--case",
        case,
        "--control",
        control,
        "--output-dir",
        str(output_dir),
        "--report-author",
        str(report_branding.get("report_author") or DEFAULT_REPORT_AUTHOR),
        "--report-entity-name",
        str(report_branding.get("report_entity_name") or ""),
        "--report-entity-logo-path",
        str(report_branding.get("report_entity_logo_path") or ""),
        "--report-title",
        "RNA-seq Python Report",
        "--report-subtitle",
        " | ".join(
            [
                item
                for item in (
                    _string_value(design_row.get("treatment1")) + " vs " + _string_value(design_row.get("treatment2"))
                    if _string_value(design_row.get("treatment1")) and _string_value(design_row.get("treatment2"))
                    else "",
                    _string_value(design_row.get("organism")),
                    _string_value(design_row.get("molecule")),
                    _string_value(design_row.get("cell")),
                )
                if item
            ]
        ),
        "--report-organism",
        _string_value(design_row.get("organism")),
        "--report-molecule",
        _string_value(design_row.get("molecule")),
        "--report-cell-context",
        _string_value(design_row.get("cell")),
        "--report-condition-description",
        _string_value(design_row.get("condition")),
        "--report-mode-label",
        "Python workflow",
        "--docx",
    ]
    process = subprocess.Popen(
        command,
        stdout=subprocess.PIPE,
        stderr=subprocess.STDOUT,
        text=True,
        bufsize=1,
    )
    output_lines: list[str] = []
    if process.stdout:
        for line in iter(process.stdout.readline, ""):
            if not line:
                break
            output_lines.append(line.rstrip())

    process.wait()
    if process.returncode != 0:
        raise RuntimeError("\n".join(output_lines[-20:]) or f"Python pipeline terminó con código {process.returncode}")

    _sync_python_report_artifacts(output_dir, design_id)
    return output_lines


def _log_analysis_event(
    kind: str,
    *,
    actor_username: str,
    analysis_type: str,
    design_id: str,
    project_name: str,
    project_owner_username: str,
    title: str,
    description: str,
) -> None:
    log_analysis_dashboard_event(
        kind,
        actor_username=actor_username,
        analysis_type=analysis_type,
        description=description,
        design_id=design_id,
        project_name=project_name,
        project_owner_username=project_owner_username,
        title=title,
    )


def _run_r_markdown(
    *,
    design_dir: Path,
    design_id: str,
    project_dir: Path,
    report_branding: dict[str, str],
    script_key: str,
    settings,
) -> subprocess.Popen[str]:
    rmd_file = settings.r_scripts_dir / f"{script_key}.Rmd"
    if not rmd_file.exists():
        raise FileNotFoundError(f"Rmd file not found for analysis_type: {script_key}")

    output_file = design_dir / f"{design_id}.html"
    command = [
        "Rscript",
        "-e",
        (
            f'rmarkdown::render("{rmd_file}", '
            f'output_file="{output_file}", '
            f'params=list(designID="{design_id}", '
            f'base_dir="{settings.project_root}", '
            f'project_dir="{project_dir}", '
            f'design_dir="{design_dir}", '
            f'report_author={json.dumps(str(report_branding.get("report_author") or DEFAULT_REPORT_AUTHOR))}, '
            f'report_entity_name={json.dumps(str(report_branding.get("report_entity_name") or ""))}, '
            f'report_entity_logo_path={json.dumps(str(report_branding.get("report_entity_logo_path") or ""))}))'
        ),
    ]
    return subprocess.Popen(
        command,
        stdout=subprocess.PIPE,
        stderr=subprocess.STDOUT,
        text=True,
        bufsize=1,
    )


def iter_analysis_events(
    project_owner_username: str,
    project_name: str,
    *,
    actor_username: str | None = None,
    preferred_variant: str | None = None,
) -> Iterator[dict[str, object]]:
    settings = get_settings()
    resolved_actor_username = actor_username or project_owner_username
    try:
        project_dir = get_project_dir(project_owner_username, project_name)
    except ValueError as exc:
        yield {
            "type": "run_failed",
            "message": str(exc),
            "project_name": project_name,
            "timestamp": _timestamp(),
        }
        return

    if not project_dir.exists():
        yield {
            "type": "run_failed",
            "message": "Proyecto no encontrado",
            "project_name": project_name,
            "timestamp": _timestamp(),
        }
        return

    try:
        excel_file = resolve_template_file(project_dir)
    except FileNotFoundError:
        yield {
            "type": "run_failed",
            "message": "Archivo template.xlsx no encontrado",
            "project_name": project_name,
            "timestamp": _timestamp(),
        }
        return

    try:
        workbook = load_workbook(excel_file)
    except Exception as exc:
        yield {
            "type": "run_failed",
            "message": f"No se pudo leer el Excel del proyecto: {exc}",
            "project_name": project_name,
            "timestamp": _timestamp(),
        }
        return

    if "design" not in workbook.sheetnames:
        yield {
            "type": "run_failed",
            "message": "Hoja 'design' no encontrada en template.xlsx",
            "project_name": project_name,
            "timestamp": _timestamp(),
        }
        return

    sheet = workbook["design"]
    try:
        headers = [cell for cell in next(sheet.iter_rows(values_only=True))]
    except StopIteration:
        yield {
            "type": "run_failed",
            "message": "La hoja 'design' está vacía",
            "project_name": project_name,
            "timestamp": _timestamp(),
        }
        return

    rows = list(sheet.iter_rows(min_row=2, values_only=True))
    total_designs = sum(
        1
        for row in rows
        if dict(zip(headers, row)).get("designID") and dict(zip(headers, row)).get("analysis_type")
    )
    processed_designs = 0

    yield {
        "type": "run_started",
        "project_name": project_name,
        "timestamp": _timestamp(),
        "total_designs": total_designs,
    }

    if total_designs == 0:
        yield {
            "type": "log",
            "message": "No se encontraron filas válidas en la hoja 'design' para ejecutar el informe.",
            "level": "warning",
            "timestamp": _timestamp(),
            "total_designs": 0,
        }
        yield {
            "type": "run_completed",
            "project_name": project_name,
            "timestamp": _timestamp(),
            "total_designs": 0,
            "processed_designs": 0,
        }
        return

    for row_index, row in enumerate(rows, start=1):
        record = dict(zip(headers, row))
        design_id = record.get("designID")
        analysis_type = record.get("analysis_type")
        if not design_id or not analysis_type:
            yield {
                "type": "log",
                "message": "Se omite una fila de la hoja 'design' porque no incluye designID o analysis_type.",
                "level": "warning",
                "timestamp": _timestamp(),
                "current_index": row_index,
                "total_designs": total_designs,
            }
            continue

        design_started_at = datetime.now(UTC)
        design_id_str = str(design_id)
        analysis_type_str = str(analysis_type)
        script_key = resolve_analysis_script_key(
            project_dir,
            analysis_type_str,
            preferred_variant=preferred_variant,
        )
        report_branding = _resolve_report_branding(project_owner_username, project_name)
        design_dir = resolve_design_output_dir(project_dir, design_id_str, script_key)
        design_dir.mkdir(parents=True, exist_ok=True)
        reset_output_dir(design_dir)

        _log_analysis_event(
            "analysis_started",
            actor_username=resolved_actor_username,
            analysis_type=script_key,
            description=f"Se ha iniciado la ejecución de {script_key} para {design_id}.",
            design_id=design_id_str,
            project_name=project_name,
            project_owner_username=project_owner_username,
            title=f"Análisis iniciado en {project_name}",
        )
        yield {
            "type": "design_started",
            "analysis_type": script_key,
            "current_index": processed_designs + 1,
            "design_id": design_id_str,
            "message": f"Running analysis for designID {design_id} using {script_key}",
            "timestamp": _timestamp(),
            "total_designs": total_designs,
        }

        exit_code = 0
        try:
            if script_key == "rna-seq-python":
                metadata_rows = _worksheet_records(workbook, "metadata")
                filtered_metadata = _filter_metadata_records(metadata_rows, record)
                output_lines = _run_python_analysis(
                    project_dir=project_dir,
                    output_dir=design_dir,
                    design_id=design_id_str,
                    design_row=record,
                    metadata_rows=filtered_metadata,
                    report_branding=report_branding,
                )
                for line in output_lines:
                    yield {
                        "type": "log",
                        "message": line.strip(),
                        "level": "info",
                        "timestamp": _timestamp(),
                        "analysis_type": script_key,
                        "current_index": processed_designs + 1,
                        "design_id": design_id_str,
                        "total_designs": total_designs,
                    }
            else:
                process = _run_r_markdown(
                    design_dir=design_dir,
                    design_id=design_id_str,
                    project_dir=project_dir,
                    report_branding=report_branding,
                    script_key=script_key,
                    settings=settings,
                )
                if process.stdout:
                    for line in iter(process.stdout.readline, ""):
                        if not line:
                            break
                        yield {
                            "type": "log",
                            "message": line.strip(),
                            "level": "info",
                            "timestamp": _timestamp(),
                            "analysis_type": script_key,
                            "current_index": processed_designs + 1,
                            "design_id": design_id_str,
                            "total_designs": total_designs,
                        }
                process.wait()
                exit_code = process.returncode
        except FileNotFoundError as exc:
            message = "Rscript no está disponible en el entorno actual" if str(exc).strip() == "[Errno 2] No such file or directory: 'Rscript'" else str(exc)
            yield {
                "type": "design_failed",
                "analysis_type": script_key,
                "current_index": processed_designs + 1,
                "design_id": design_id_str,
                "message": message,
                "timestamp": _timestamp(),
                "total_designs": total_designs,
            }
            processed_designs += 1
            continue
        except Exception as exc:
            yield {
                "type": "design_failed",
                "analysis_type": script_key,
                "current_index": processed_designs + 1,
                "design_id": design_id_str,
                "message": f"No se pudo iniciar o completar el análisis: {exc}",
                "timestamp": _timestamp(),
                "total_designs": total_designs,
            }
            processed_designs += 1
            continue

        duration_seconds = max(0.0, (datetime.now(UTC) - design_started_at).total_seconds())
        if exit_code == 0:
            try:
                _apply_docx_branding(
                    design_dir / f"{design_id_str}.docx",
                    design_dir,
                    report_branding,
                    settings,
                )
            except Exception as exc:
                yield {
                    "type": "log",
                    "message": f"WARNING: No se pudo aplicar branding DOCX: {exc}",
                    "level": "warning",
                    "timestamp": _timestamp(),
                    "analysis_type": script_key,
                    "current_index": processed_designs + 1,
                    "design_id": design_id_str,
                    "total_designs": total_designs,
                }
            _log_analysis_event(
                "analysis_completed",
                actor_username=resolved_actor_username,
                analysis_type=script_key,
                description=f"El análisis {script_key} para {design_id} terminó correctamente.",
                design_id=design_id_str,
                project_name=project_name,
                project_owner_username=project_owner_username,
                title=f"Análisis completado en {project_name}",
            )
            yield {
                "type": "design_completed",
                "analysis_type": script_key,
                "current_index": processed_designs + 1,
                "design_id": design_id_str,
                "duration_seconds": duration_seconds,
                "message": f"Finished analysis for designID {design_id}",
                "timestamp": _timestamp(),
                "total_designs": total_designs,
            }
        else:
            _log_analysis_event(
                "analysis_failed",
                actor_username=resolved_actor_username,
                analysis_type=script_key,
                description=f"El análisis {script_key} para {design_id} terminó con código {exit_code}.",
                design_id=design_id_str,
                project_name=project_name,
                project_owner_username=project_owner_username,
                title=f"Análisis con incidencias en {project_name}",
            )
            yield {
                "type": "design_failed",
                "analysis_type": script_key,
                "current_index": processed_designs + 1,
                "design_id": design_id_str,
                "duration_seconds": duration_seconds,
                "exit_code": exit_code,
                "message": f"El análisis de {design_id} terminó con código {exit_code}",
                "timestamp": _timestamp(),
                "total_designs": total_designs,
            }

        if design_dir.exists():
            try:
                clean_resultados(design_dir, design_id_str)
                yield {
                    "type": "cleanup_completed",
                    "analysis_type": script_key,
                    "current_index": processed_designs + 1,
                    "design_id": design_id_str,
                    "message": f"Cleaned {design_id} folder, kept report artifacts",
                    "timestamp": _timestamp(),
                    "total_designs": total_designs,
                }
            except Exception as exc:
                yield {
                    "type": "cleanup_failed",
                    "analysis_type": script_key,
                    "current_index": processed_designs + 1,
                    "design_id": design_id_str,
                    "message": f"WARNING: Could not clean Resultados folder: {exc}",
                    "timestamp": _timestamp(),
                    "total_designs": total_designs,
                }

        processed_designs += 1

    yield {
        "type": "run_completed",
        "project_name": project_name,
        "timestamp": _timestamp(),
        "total_designs": total_designs,
        "processed_designs": processed_designs,
    }


def stream_analysis(username: str, project_name: str) -> Iterator[str]:
    for payload in iter_analysis_events(username, project_name, actor_username=username):
        yield _stream_event(payload)
    yield "data:---FIN---\n\n"
